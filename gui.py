# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2025 therudywolf <https://github.com/therudywolf>
#
# This file is part of ForestOptiLM / Nocturne Data Forge.
# ForestOptiLM is free software: you can redistribute it and/or modify it
# under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# ForestOptiLM is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY
# or FITNESS FOR A PARTICULAR PURPOSE. See the GNU Affero General Public
# License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with ForestOptiLM. If not, see <https://www.gnu.org/licenses/>.
"""
Nocturne Data Forge — GUI (CustomTkinter), тёмная тема.
Всегда: batch Map-Reduce по всем файлам / папке.
"""
from __future__ import annotations

import asyncio
import datetime as dt
import logging
import os
import queue
import sys
import threading
import tkinter
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from typing import Any, Final, Literal

import customtkinter as ctk
import httpx
import pandas as pd

import lm_studio_api as lmsapi
from lmstudio_config import (
    get_default_model_optional,
    lmstudio_root_url,
    load_ui_runtime_state,
    sanitize_for_log,
    save_ui_runtime_state,
    validate_lmstudio_url,
)
from parser import ParseError, compute_dynamic_chunk_size, parse_file
from processor import (
    API_BASE,
    API_KEY,
    CONTEXT_FALLBACK,
    SYSTEM_PROMPT_MAP,
    check_vision_capability,
    compute_job_id,
    fetch_models_info,
    categorize_models,
    answer_with_context,
    set_runtime_modes,
    set_runtime_limits,
    summarize_model_tokens_by_category,
    resolve_runtime_model_context,
    run_batching,
    run_map_reduce,
    test_lmstudio_connection,
)
from pipeline import build_index, query_index
from run_config import RunConfig
from notebook_gui import NotebookUIMixin
import connection_presets as cp

logger = logging.getLogger("nocturne")

MSG_PROGRESS  = "progress"
MSG_RESULT    = "result"
MSG_RESULT_DF = "result_df"
MSG_ERROR     = "error"
MSG_TRACE     = "trace"
MSG_JOB_ID    = "job_id"
MAX_UI_WORKERS = 4

FILE_TYPES = [
    ("Документы / Код",
     "*.txt *.md *.log *.ini *.rtf *.pdf *.docx *.odt *.epub "
     "*.html *.htm *.py *.js *.ts *.tsx *.c *.cpp *.h *.java "
     "*.go *.rs *.kt *.sql *.sh *.bat *.ps1 *.toml *.cfg *.properties"),
    ("Изображения (vision)", "*.png *.jpg *.jpeg *.webp *.gif *.bmp *.tif *.tiff"),
    ("Таблицы",  "*.csv *.xlsx *.xls *.json *.yaml *.yml"),
    ("Аудио (транскрипция)", "*.mp3 *.wav *.m4a *.ogg *.flac *.opus *.aac *.wma"),
    ("Архивы",   "*.zip *.tar *.tar.gz *.tgz *.gz"),
    ("Все файлы", "*.*"),
]


# Вкладки. NotebookLM-режим («Блокноты») идёт первым и по умолчанию — это лицо
# приложения; массовый Map-Reduce-анализ — вторичен.
TAB_NOTEBOOKS: Final = "📓  Блокноты"
TAB_RESULT: Final = "📊  Результат"
TAB_LOGS: Final = "🧾  Логи"
TAB_RAG: Final = "🔎  RAG"

import md3 as _md3  # noqa: E402

# Material Design 3 status colours (used for transient status-line messages).
_STATUS_OK = _md3.SUCCESS
_STATUS_WARN = _md3.WARNING
_STATUS_ERR = _md3.ERROR


def _apply_brand_theme() -> None:
    """Apply the Material Design 3 dark theme globally (colour roles, shape,
    surfaces) — see md3.apply()."""
    try:
        _md3.apply()
    except Exception:
        pass


class NocturneApp(NotebookUIMixin, ctk.CTk):
    def __init__(self) -> None:
        super().__init__()
        self.title("Nocturne Data Forge")
        self.geometry("1280x840")
        self.minsize(1080, 720)
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        _apply_brand_theme()

        self._file_path: Path | None = None
        self._folder_path: Path | None = None
        self._last_result_text: str = ""
        self._last_result_df: pd.DataFrame | None = None
        self._result_type: Literal["text", "table"] = "text"
        self._queue: queue.Queue[dict[str, Any]] = queue.Queue()
        self._running = False
        self._stop_requested = False
        self._active_job_id: str | None = None
        self._log_entries: list[tuple[str, str]] = []  # (phase, formatted_line)
        self._pending_log_lines: list[str] = []
        self._log_flush_scheduled = False
        self._log_flush_after_id: str | None = None
        self._log_filter_var = ctk.StringVar(value="all")
        # context_length per model, populated after "Обновить модели"
        self._model_ctx: dict[str, int] = {}
        self._model_ctx_source: dict[str, str] = {}
        self._models_by_kind: dict[str, list[str]] = {
            "chat": [], "vision": [], "embedding": [], "reasoning": [],
        }
        self._cfg_default_model = get_default_model_optional()
        self._runtime_state = load_ui_runtime_state()
        self._runtime_state_ready = False
        self._model_poll_after_id: str | None = None
        self._queue_poll_after_id: str | None = None
        self._model_poll_active = False
        self._closing = False

        self._build_ui()
        self._set_window_icon()
        self._apply_runtime_state_to_widgets()
        self._runtime_state_ready = True
        self._bind_runtime_state_watchers()
        self.protocol("WM_DELETE_WINDOW", self._on_close)
        self.after(300, self._on_fetch_models)
        self.after(600, self._maybe_first_run_wizard)
        # Открыть конкретную вкладку на старте (отладка/удобство): NOCTURNE_STARTUP_TAB=Блокноты
        _startup_tab = os.environ.get("NOCTURNE_STARTUP_TAB", "").strip()
        if _startup_tab:
            self.after(250, lambda t=_startup_tab: self._select_tab_by_hint(t))

    # ------------------------------------------------------------------ #
    #  Layout
    # ------------------------------------------------------------------ #

    def _build_ui(self) -> None:
        # ---- sidebar ----
        sidebar_container = ctk.CTkFrame(self, width=300, corner_radius=0,
                                         fg_color=_md3.SURFACE_CONTAINER_HIGH)
        sidebar_container.pack(side="left", fill="y")
        sidebar_container.pack_propagate(False)
        sidebar = ctk.CTkScrollableFrame(
            sidebar_container,
            corner_radius=0,
            fg_color="transparent",
            width=290,
        )
        sidebar.pack(fill="both", expand=True)
        self._build_sidebar(sidebar)

        # ---- main panel ----
        main = ctk.CTkFrame(self, corner_radius=0, fg_color="transparent")
        main.pack(side="left", fill="both", expand=True, padx=18, pady=18)
        self._build_main(main)

    def _build_sidebar(self, sb: ctk.CTkFrame) -> None:
        pad = {"padx": 14}

        ctk.CTkLabel(sb, text="Nocturne Data Forge",
                     font=ctk.CTkFont(size=15, weight="bold")
                     ).pack(anchor="w", pady=(14, 0), **pad)
        ctk.CTkLabel(sb, text="Локальный анализ данных через LLM",
                     text_color=_md3.ON_SURFACE_VARIANT, font=ctk.CTkFont(size=11)
                     ).pack(anchor="w", pady=(0, 10), **pad)

        # Провайдер — быстрая настройка сервера (LM Studio / Ollama / OpenAI-совм.)
        ctk.CTkLabel(sb, text="Провайдер (сервер LLM)").pack(anchor="w", pady=(8, 0), **pad)
        self._provider_var = ctk.StringVar(
            value=(cp.get_preset(cp.DEFAULT_PRESET_KEY) or cp.PRESETS[0]).label
        )
        self._provider_menu = ctk.CTkOptionMenu(
            sb, variable=self._provider_var, values=cp.preset_labels(),
            dynamic_resizing=False, width=250, command=self._on_provider_change,
        )
        self._provider_menu.pack(fill="x", pady=3, **pad)
        self._provider_hint = ctk.CTkLabel(
            sb, text="", text_color=_md3.ON_SURFACE_VARIANT, font=ctk.CTkFont(size=11),
            wraplength=250, justify="left",
        )
        self._provider_hint.pack(anchor="w", pady=(0, 4), **pad)

        # API Base URL
        ctk.CTkLabel(sb, text="API Base URL").pack(anchor="w", pady=(8, 0), **pad)
        self._url_var = ctk.StringVar(value=API_BASE)
        ctk.CTkEntry(sb, textvariable=self._url_var).pack(fill="x", pady=3, **pad)

        # API Key
        ctk.CTkLabel(sb, text="API Key (если требуется)").pack(anchor="w", pady=(6, 0), **pad)
        self._api_key_var = ctk.StringVar(value=API_KEY)
        ctk.CTkEntry(sb, textvariable=self._api_key_var, show="*"
                     ).pack(fill="x", pady=3, **pad)

        # Buttons row
        ctk.CTkButton(sb, text="Обновить модели", height=30,
                      command=self._on_fetch_models
                      ).pack(fill="x", pady=(10, 3), **pad)
        ctk.CTkButton(sb, text="Проверить подключение", height=28,
                      fg_color="transparent", border_width=1,
                      command=self._on_test_connection
                      ).pack(fill="x", pady=(0, 3), **pad)
        ctk.CTkButton(sb, text="Скачать модель…", height=28,
                      fg_color="transparent", border_width=1,
                      command=self._on_download_model_dialog
                      ).pack(fill="x", pady=(0, 8), **pad)

        # LLM model
        ctk.CTkLabel(sb, text="LLM модель (для ответов)").pack(anchor="w", pady=(6, 0), **pad)
        self._model_var = ctk.StringVar(value="")
        self._model_menu = ctk.CTkOptionMenu(
            sb, variable=self._model_var,
            values=["(нажмите Обновить модели)"],
            dynamic_resizing=False, width=250,
        )
        self._model_menu.pack(fill="x", pady=3, **pad)
        ctk.CTkLabel(sb, text="Чат-модель для генерации ответов.",
                     text_color=_md3.ON_SURFACE_VARIANT, font=ctk.CTkFont(size=11),
                     wraplength=250, justify="left",
                     ).pack(anchor="w", pady=(0, 4), **pad)
        ctk.CTkButton(
            sb,
            text="Пересчитать контекст модели",
            height=28,
            fg_color="transparent",
            border_width=1,
            command=self._on_recalc_selected_model_context,
        ).pack(fill="x", pady=(0, 8), **pad)

        # Separator
        # Context length — auto-detected label
        self._ctx_label = ctk.CTkLabel(
            sb, text="Контекст модели: не определён",
            text_color=_md3.ON_SURFACE_VARIANT, font=ctk.CTkFont(size=11),
            wraplength=250, justify="left",
        )
        self._ctx_label.pack(anchor="w", pady=(8, 4), **pad)

        ctk.CTkLabel(
            sb,
            text="Контекст берётся автоматически из LM Studio "
                 "(loaded_context_length/context_length).",
            text_color=_md3.ON_SURFACE_VARIANT,
            font=ctk.CTkFont(size=11),
            wraplength=250,
            justify="left",
        ).pack(anchor="w", pady=(0, 8), **pad)

        # ── Модели поиска (нужны и «Блокнотам», и Map-Reduce) ───────────
        ctk.CTkLabel(sb, text="Embedding-модель (поиск/RAG)", font=ctk.CTkFont(size=12, weight="bold")
                     ).pack(anchor="w", pady=(8, 0), **pad)
        self._embedding_model_var = ctk.StringVar(value="")
        self._embedding_menu = ctk.CTkOptionMenu(
            sb,
            variable=self._embedding_model_var,
            values=["(нажмите Обновить модели)"],
            dynamic_resizing=False,
            width=250,
        )
        self._embedding_menu.pack(fill="x", pady=(3, 8), **pad)

        ctk.CTkLabel(sb, text="Vision-модель (описание картинок)", font=ctk.CTkFont(size=12, weight="bold")
                     ).pack(anchor="w", pady=(6, 0), **pad)
        self._vision_model_var = ctk.StringVar(value="")
        self._vision_menu = ctk.CTkOptionMenu(
            sb, variable=self._vision_model_var,
            values=["(нажмите Обновить модели)"],
            dynamic_resizing=False, width=250,
        )
        self._vision_menu.pack(fill="x", pady=3, **pad)
        ctk.CTkButton(
            sb, text="Проверить Vision",
            height=28,
            fg_color="transparent",
            border_width=1,
            command=self._on_test_vision,
        ).pack(fill="x", pady=(0, 8), **pad)

        ctk.CTkLabel(sb, text="Режим API").pack(anchor="w", pady=(4, 0), **pad)
        self._api_mode_var = ctk.StringVar(value="native")
        self._api_mode_menu = ctk.CTkOptionMenu(
            sb, variable=self._api_mode_var,
            values=["native", "openai"],
            dynamic_resizing=False, width=250,
            command=lambda _v: self._on_runtime_mode_changed(),
        )
        self._api_mode_menu.pack(fill="x", pady=(3, 8), **pad)

        # ── Массовая обработка файлов (скрывается в режиме «Блокноты») ───
        # Все Map-Reduce-контролы (воркеры/composer/scout/лимиты/профили)
        # собраны в один фрейм: в NotebookLM-режиме они не нужны и только шумят,
        # поэтому _on_tab_changed прячет весь блок целиком.
        self._sb_mapreduce = ctk.CTkFrame(sb, fg_color="transparent")
        mr = self._sb_mapreduce
        mr.pack(fill="x")
        ctk.CTkLabel(mr, text="⚙️  Массовая обработка файлов",
                     font=ctk.CTkFont(size=12, weight="bold"), text_color=_md3.PRIMARY,
                     ).pack(anchor="w", pady=(8, 2), **pad)

        ctk.CTkLabel(mr, text="Параллельных воркеров (1–4)",
                     font=ctk.CTkFont(size=12, weight="bold")
                     ).pack(anchor="w", pady=(6, 0), **pad)
        ctk.CTkLabel(
            mr, text="Сколько файлов/чанков обрабатывать одновременно.",
            text_color=_md3.ON_SURFACE_VARIANT, font=ctk.CTkFont(size=11),
            wraplength=250, justify="left",
        ).pack(anchor="w", pady=(0, 2), **pad)
        self._workers_var = ctk.StringVar(value="3")
        self._workers_seg = ctk.CTkSegmentedButton(
            mr,
            values=["1", "2", "3", "4"],
            variable=self._workers_var,
            dynamic_resizing=False,
        )
        self._workers_seg.pack(fill="x", pady=(0, 8), **pad)

        self._composer_use_var = ctk.BooleanVar(value=False)
        self._composer_check = ctk.CTkCheckBox(
            mr,
            text="Отдельная модель для сборки итога (reduce)",
            variable=self._composer_use_var,
            command=self._on_composer_toggle,
        )
        self._composer_check.pack(anchor="w", pady=(4, 2), **pad)
        self._composer_model_var = ctk.StringVar(value="")
        self._composer_menu = ctk.CTkOptionMenu(
            mr, variable=self._composer_model_var,
            values=["(нажмите Обновить модели)"],
            dynamic_resizing=False, width=250,
        )
        self._composer_menu.pack(fill="x", pady=(0, 8), **pad)
        self._composer_menu.configure(state="disabled")

        self._low_vram_var = ctk.BooleanVar(value=True)
        self._low_vram_check = ctk.CTkCheckBox(
            mr,
            text="Low VRAM Sequential",
            variable=self._low_vram_var,
            command=self._on_runtime_mode_changed,
        )
        self._low_vram_check.pack(anchor="w", pady=(2, 2), **pad)
        self._advanced_visible = ctk.BooleanVar(value=False)
        ctk.CTkCheckBox(
            mr,
            text="Дополнительно (лимиты, scout-модель)",
            variable=self._advanced_visible,
            command=self._toggle_advanced_panel,
        ).pack(anchor="w", pady=(10, 2), **pad)
        self._advanced_frame = ctk.CTkFrame(mr, fg_color="transparent")
        self._max_reduce_tokens_var = ctk.StringVar(value=str(self._runtime_state.get("max_reduce_input_tokens", 24000)))
        self._max_chunk_tokens_var = ctk.StringVar(value=str(self._runtime_state.get("max_chunk_tokens", 6000)))
        ctk.CTkLabel(
            self._advanced_frame, text="Лимиты токенов", font=ctk.CTkFont(size=12, weight="bold"),
        ).pack(anchor="w", pady=(4, 0), **pad)
        ctk.CTkLabel(self._advanced_frame, text="MAX_REDUCE_INPUT_TOKENS").pack(anchor="w", pady=(4, 0), **pad)
        ctk.CTkEntry(self._advanced_frame, textvariable=self._max_reduce_tokens_var).pack(
            fill="x", pady=(2, 4), **pad,
        )
        ctk.CTkLabel(self._advanced_frame, text="NOCTURNE_MAX_CHUNK_TOKENS").pack(anchor="w", pady=(2, 0), **pad)
        ctk.CTkEntry(self._advanced_frame, textvariable=self._max_chunk_tokens_var).pack(
            fill="x", pady=(2, 4), **pad,
        )

        ctk.CTkLabel(mr, text="Большой корпус", font=ctk.CTkFont(size=12, weight="bold")
                     ).pack(anchor="w", pady=(8, 0), **pad)
        self._scout_var = ctk.BooleanVar(value=bool(self._runtime_state.get("scout_mode", False)))
        self._scout_check = ctk.CTkCheckBox(
            mr,
            text="Scout-pass (быстрая релевантность)",
            variable=self._scout_var,
            command=self._persist_runtime_state,
        )
        self._scout_check.pack(anchor="w", pady=(2, 2), **pad)
        self._scout_threshold_var = ctk.StringVar(
            value=str(self._runtime_state.get("scout_threshold", 0.35))
        )
        self._scout_model_var = ctk.StringVar(value=str(self._runtime_state.get("selected_scout_model", "")))
        ctk.CTkLabel(self._advanced_frame, text="Порог релевантности (0–1)").pack(
            anchor="w", pady=(2, 0), **pad,
        )
        ctk.CTkEntry(self._advanced_frame, textvariable=self._scout_threshold_var, width=80).pack(
            fill="x", pady=(2, 4), **pad,
        )
        ctk.CTkLabel(self._advanced_frame, text="Scout-модель (пусто = MAP-модель)").pack(
            anchor="w", pady=(2, 0), **pad,
        )
        self._scout_menu = ctk.CTkOptionMenu(
            self._advanced_frame,
            variable=self._scout_model_var,
            values=["(как MAP-модель)"],
            dynamic_resizing=False,
            width=250,
        )
        self._scout_menu.pack(fill="x", pady=(2, 4), **pad)
        mode_row = ctk.CTkFrame(mr, fg_color="transparent")
        mode_row.pack(fill="x", pady=(2, 4), **pad)
        ctk.CTkButton(
            mode_row, text="Быстро", width=72, height=26,
            command=lambda: self._apply_run_profile("quick_scan"),
        ).pack(side="left", padx=(0, 4))
        ctk.CTkButton(
            mode_row, text="Глубоко", width=72, height=26,
            command=lambda: self._apply_run_profile("deep_audit"),
        ).pack(side="left", padx=(0, 4))
        ctk.CTkButton(
            mode_row, text="1M+", width=56, height=26,
            fg_color=_md3.PRIMARY_CONTAINER,
            command=lambda: self._apply_run_profile("large_corpus"),
        ).pack(side="left")
        ctk.CTkButton(
            mr,
            text="Пресет: корпус 1M+",
            height=28,
            fg_color=_md3.PRIMARY_CONTAINER,
            hover_color=_md3.PRIMARY_CONTAINER_HOVER,
            command=self._on_large_corpus_preset,
        ).pack(fill="x", pady=(2, 4), **pad)

    def _build_main(self, main: ctk.CTkFrame) -> None:
        # Контролы массового Map-Reduce-анализа собраны в один фрейм — его прячем,
        # когда активна вкладка «Блокноты» (NotebookLM-режим), чтобы не загромождать.
        self._analysis_panel = ctk.CTkFrame(main, fg_color="transparent")
        self._analysis_panel.pack(fill="x")
        panel = self._analysis_panel

        # File / folder row
        btn_row = ctk.CTkFrame(panel, fg_color="transparent")
        btn_row.pack(anchor="w", fill="x", pady=(0, 4))
        ctk.CTkButton(btn_row, text="Выбрать файл", width=140,
                      command=self._on_select_file).pack(side="left", padx=(0, 8))
        ctk.CTkButton(btn_row, text="Выбрать папку", width=140,
                      command=self._on_select_folder).pack(side="left")

        self._file_label = ctk.CTkLabel(
            panel, text="Файл или папка не выбраны",
            text_color=_md3.ON_SURFACE_VARIANT, anchor="w",
        )
        self._file_label.pack(anchor="w", pady=(2, 10), fill="x")

        # Query
        ctk.CTkLabel(panel, text="Запрос — что сделать с файлами?"
                     ).pack(anchor="w", pady=(0, 2))
        self._query_text = ctk.CTkTextbox(panel, height=88, wrap="word")
        self._query_text.pack(fill="x", pady=(0, 10))

        # Action buttons
        act_row = ctk.CTkFrame(panel, fg_color="transparent")
        act_row.pack(anchor="w", fill="x", pady=(0, 8))
        self._start_btn = ctk.CTkButton(
            act_row, text="▶  СТАРТ", width=120,
            fg_color=_md3.PRIMARY_CONTAINER, hover_color=_md3.PRIMARY_CONTAINER_HOVER,
            command=self._on_start,
        )
        self._start_btn.pack(side="left", padx=(0, 8))
        self._stop_btn = ctk.CTkButton(
            act_row, text="■  Стоп", width=90,
            fg_color=_md3.ERROR_CONTAINER, hover_color="#A8000C",
            state="disabled", command=self._on_stop,
        )
        self._stop_btn.pack(side="left")
        ctk.CTkButton(
            act_row, text="Оценка (dry-run)", width=130,
            fg_color="transparent",
            border_width=1,
            command=self._on_dry_run,
        ).pack(side="left", padx=(8, 0))
        ctk.CTkButton(
            act_row, text="История", width=80,
            fg_color="transparent",
            border_width=1,
            command=self._on_run_history,
        ).pack(side="left", padx=(6, 0))
        ctk.CTkButton(
            act_row, text="Продолжить", width=100,
            fg_color=_md3.PRIMARY_CONTAINER,
            hover_color=_md3.PRIMARY_CONTAINER_HOVER,
            command=self._on_resume_job,
        ).pack(side="left", padx=(6, 0))

        # Progress
        self._progress_bar = ctk.CTkProgressBar(panel)
        self._progress_bar.pack(fill="x", pady=(0, 2))
        self._progress_bar.set(0)
        self._status_label = ctk.CTkLabel(
            panel, text="Готов к работе", anchor="w", text_color=_md3.ON_SURFACE_VARIANT,
        )
        self._status_label.pack(anchor="w", pady=(0, 4))
        self._preflight_label = ctk.CTkLabel(
            panel,
            text="Preflight: выберите файл и модель",
            anchor="w",
            text_color=_md3.ON_SURFACE_VARIANT,
            wraplength=820,
            justify="left",
        )
        self._preflight_label.pack(anchor="w", pady=(0, 8))
        self._loaded_model_label = ctk.CTkLabel(
            panel,
            text="Активная модель в LM Studio: неизвестно",
            anchor="w",
            text_color=_md3.ON_SURFACE_VARIANT,
        )
        self._loaded_model_label.pack(anchor="w", pady=(0, 8))

        # Meta-prompt preview label (shown when composer generates a directive)
        self._meta_prompt_label = ctk.CTkLabel(
            panel,
            text="",
            anchor="w",
            text_color=_STATUS_OK,
            wraplength=820,
            justify="left",
        )
        self._meta_prompt_label.pack(anchor="w", pady=(0, 4))
        self._meta_prompt_label.pack_forget()  # hidden until first meta_plan_done

        # Вкладки: NotebookLM («Блокноты») — первая и по умолчанию.
        self._tabs = ctk.CTkTabview(main, command=self._on_tab_changed)
        self._tabs.pack(fill="both", expand=True, pady=(0, 6))
        self._notebooks_tab = self._tabs.add(TAB_NOTEBOOKS)
        self._result_tab = self._tabs.add(TAB_RESULT)
        self._logs_tab = self._tabs.add(TAB_LOGS)
        self._rag_tab = self._tabs.add(TAB_RAG)
        self._tabs.set(TAB_NOTEBOOKS)
        try:
            self._tabs._segmented_button.configure(
                font=ctk.CTkFont(size=14, weight="bold"))
        except Exception:
            pass

        result_toolbar = ctk.CTkFrame(self._result_tab, fg_color="transparent")
        result_toolbar.pack(fill="x", pady=(0, 6))
        ctk.CTkButton(result_toolbar, text="Сохранить…", width=100,
                      command=self._on_save_result).pack(side="left", padx=(0, 8))
        ctk.CTkButton(result_toolbar, text="Экспорт MD", width=100,
                      command=lambda: self._export_result(".md")).pack(side="left", padx=(0, 8))
        ctk.CTkButton(result_toolbar, text="Экспорт JSON evidence", width=150,
                      command=self._on_export_evidence_json).pack(side="left", padx=(0, 8))
        ctk.CTkButton(result_toolbar, text="История запусков", width=130,
                      command=self._on_show_run_history).pack(side="right")

        self._result_text = ctk.CTkTextbox(self._result_tab, wrap="word")
        self._result_text.pack(fill="both", expand=True)
        self._result_text.bind("<Key>", self._on_result_keypress)

        log_toolbar = ctk.CTkFrame(self._logs_tab, fg_color="transparent")
        log_toolbar.pack(fill="x", pady=(0, 6))
        ctk.CTkButton(log_toolbar, text="Очистить", width=90,
                      command=self._on_clear_logs).pack(side="left", padx=(0, 8))
        ctk.CTkButton(log_toolbar, text="Копировать", width=90,
                      command=self._on_copy_logs).pack(side="left", padx=(0, 8))
        ctk.CTkButton(log_toolbar, text="Сохранить…", width=100,
                      command=self._on_save_logs).pack(side="left")
        ctk.CTkLabel(log_toolbar, text="Фильтр:").pack(side="right", padx=(8, 4))
        self._log_filter_combo = ctk.CTkComboBox(
            log_toolbar,
            values=[
                "all", "trace", "preflight", "extract", "map", "map_metrics", "vision_map",
                "reduce", "reduce_merge", "retry", "summary", "quality_metrics", "batch",
                "error", "general", "section_reduce", "synthesize",
            ],
            variable=self._log_filter_var,
            width=160,
            command=self._on_log_filter_change,
        )
        self._log_filter_combo.pack(side="right")
        self._log_text = ctk.CTkTextbox(self._logs_tab, wrap="word")
        self._log_text.pack(fill="both", expand=True)
        self._build_rag_tab(self._rag_tab)
        self._build_notebooks_tab(self._notebooks_tab)

        # Save row (только для результата Map-Reduce — прячем на «Блокнотах»)
        self._save_row = ctk.CTkFrame(main, fg_color="transparent")
        save_row = self._save_row
        save_row.pack(anchor="w", fill="x")
        ctk.CTkButton(save_row, text="Копировать", width=110,
                      command=self._on_copy_result).pack(side="left", padx=(0, 8))
        ctk.CTkButton(save_row, text="Сохранить…", width=110,
                      command=self._on_save_result).pack(side="left")

        # Стартуем на вкладке «Блокноты» → сразу применяем её раскладку.
        self._on_tab_changed()

    def _select_tab_by_hint(self, hint: str) -> None:
        """Выбрать вкладку по подстроке имени (имена содержат эмодзи)."""
        low = hint.lower()
        for name in (TAB_NOTEBOOKS, TAB_RESULT, TAB_LOGS, TAB_RAG):
            if low in name.lower():
                try:
                    self._tabs.set(name)
                    self._on_tab_changed()  # .set() не зовёт command — вызываем сами
                except Exception:
                    pass
                return

    def _on_tab_changed(self, *_args: object) -> None:
        """Контекстный UI: Map-Reduce-настройки (сайдбар) и панель запуска
        показываем ТОЛЬКО в файловом режиме (Результат/Логи). В «Блокнотах» и
        «RAG» их прячем — чистый вид, без нерелевантных воркеров/composer/scout."""
        try:
            cur = self._tabs.get()
        except Exception:
            return
        file_mode = cur in (TAB_RESULT, TAB_LOGS)  # массовая обработка файлов
        try:
            if file_mode:
                if not self._sb_mapreduce.winfo_ismapped():
                    self._sb_mapreduce.pack(fill="x")
            else:
                self._sb_mapreduce.pack_forget()
        except Exception:
            pass
        if file_mode:
            if not self._analysis_panel.winfo_ismapped():
                self._analysis_panel.pack(fill="x", before=self._tabs)
            if not self._save_row.winfo_ismapped():
                self._save_row.pack(anchor="w", fill="x")
            if not self._model_poll_active and not self._closing:
                self._poll_loaded_model()
        else:
            self._analysis_panel.pack_forget()
            self._save_row.pack_forget()

    def _build_rag_tab(self, parent: ctk.CTkFrame) -> None:
        row1 = ctk.CTkFrame(parent, fg_color="transparent")
        row1.pack(fill="x", pady=(0, 8))
        ctk.CTkLabel(row1, text="Директория индекса").pack(side="left", padx=(0, 8))
        self._rag_index_dir_var = ctk.StringVar(value=str(self._runtime_state.get("rag_index_dir", ".nocturne_index")))
        ctk.CTkEntry(row1, textvariable=self._rag_index_dir_var).pack(side="left", fill="x", expand=True)

        row2 = ctk.CTkFrame(parent, fg_color="transparent")
        row2.pack(fill="x", pady=(0, 8))
        ctk.CTkLabel(row2, text="top_k").pack(side="left", padx=(0, 8))
        self._rag_top_k_var = ctk.StringVar(value=str(self._runtime_state.get("rag_top_k", 8)))
        ctk.CTkEntry(row2, textvariable=self._rag_top_k_var, width=80).pack(side="left")
        ctk.CTkButton(row2, text="Построить индекс", width=140, command=self._on_build_index).pack(side="left", padx=(12, 8))
        ctk.CTkButton(row2, text="Задать вопрос", width=120, command=self._on_rag_ask).pack(side="left")

        ctk.CTkLabel(parent, text="Вопрос к индексу").pack(anchor="w", pady=(4, 2))
        self._rag_question_text = ctk.CTkTextbox(parent, height=100, wrap="word")
        self._rag_question_text.pack(fill="x", pady=(0, 8))

        ctk.CTkLabel(parent, text="Ответ RAG").pack(anchor="w", pady=(2, 2))
        self._rag_answer_text = ctk.CTkTextbox(parent, wrap="word")
        self._rag_answer_text.pack(fill="both", expand=True)

    # ------------------------------------------------------------------ #
    #  Event handlers
    # ------------------------------------------------------------------ #

    def _on_result_keypress(self, event: Any) -> str | None:
        ctrl = bool((event.state or 0) & 0x4)
        if ctrl and event.keysym.lower() in {"c", "a"}:
            return None
        if event.keysym in {"Left", "Right", "Up", "Down",
                             "Home", "End", "Prior", "Next"}:
            return None
        return "break"

    def _on_copy_result(self) -> None:
        text = self._result_text.get("1.0", "end-1c")
        self.clipboard_clear()
        self.clipboard_append(text)
        self._set_status("Скопировано в буфер обмена")

    def _append_log_line(self, line: str, phase: str = "general") -> None:
        if not line:
            return
        line = sanitize_for_log(line)
        stamp = dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        formatted = f"{stamp} | {phase:14} | {line}"
        self._log_entries.append((phase, formatted))
        if len(self._log_entries) > 8000:
            self._log_entries = self._log_entries[-6000:]
        flt = self._log_filter_var.get()
        show = (
            flt == "all"
            or flt == phase
            or (flt == "error" and ("ERROR" in line or "[ERROR]" in line))
        )
        if show:
            self._pending_log_lines.append(formatted)
            if not self._log_flush_scheduled and not self._closing and self.winfo_exists():
                self._log_flush_scheduled = True
                self._log_flush_after_id = self.after(150, self._flush_logs_to_ui)

    def _flush_logs_to_ui(self) -> None:
        self._log_flush_scheduled = False
        self._log_flush_after_id = None
        if self._closing or not self.winfo_exists():
            return
        if not self._pending_log_lines:
            return
        chunk = "\n".join(self._pending_log_lines) + "\n"
        self._pending_log_lines.clear()
        self._log_text.insert("end", chunk)
        self._log_text.see("end")

    def _on_log_filter_change(self, _choice: str | None = None) -> None:
        flt = self._log_filter_var.get()
        self._log_text.delete("1.0", "end")
        for phase, formatted in self._log_entries:
            if (
                flt == "all"
                or flt == phase
                or (flt == "error" and "ERROR" in formatted)
            ):
                self._log_text.insert("end", formatted + "\n")
        self._log_text.see("end")

    def _on_clear_logs(self) -> None:
        self._log_entries.clear()
        self._pending_log_lines.clear()
        self._log_text.delete("1.0", "end")

    def _on_copy_logs(self) -> None:
        text = self._log_text.get("1.0", "end-1c")
        self.clipboard_clear()
        self.clipboard_append(text)
        self._set_status("Логи скопированы в буфер")

    def _on_save_logs(self) -> None:
        p = ctk.filedialog.asksaveasfilename(
            defaultextension=".log",
            filetypes=[("Log", "*.log"), ("Text", "*.txt"), ("Все", "*.*")],
        )
        if p:
            all_lines = [ln for _, ln in self._log_entries]
            Path(p).write_text("\n".join(all_lines), encoding="utf-8")
            self._set_status(f"Логи сохранены: {p}")

    def _pick_embedding_model(self) -> str:
        selected = self._embedding_model_var.get().strip()
        if selected and not selected.startswith("("):
            return selected
        candidates = [m for m in (self._model_menu.cget("values") or []) if isinstance(m, str)]
        embed = next((m for m in candidates if "embed" in m.lower()), "")
        if embed:
            self._embedding_model_var.set(embed)
            return embed
        return ""

    def _assets_dir(self) -> Path:
        """Каталог assets — и из исходников, и из собранного бандла (PyInstaller)."""
        base = Path(getattr(sys, "_MEIPASS", Path(__file__).resolve().parent))
        return base / "assets"

    def _set_window_icon(self) -> None:
        """Иконка окна: .ico на Windows + PNG-фолбэк (iconphoto) кроссплатформенно."""
        assets = self._assets_dir()
        ico = assets / "icon.ico"
        png = assets / "icon.png"
        try:
            if sys.platform == "win32" and ico.is_file():
                self.iconbitmap(default=str(ico))
        except Exception:
            pass
        try:
            if png.is_file():
                # держим ссылку, иначе Tk соберёт изображение мусором
                self._icon_photo = tkinter.PhotoImage(file=str(png))
                self.iconphoto(True, self._icon_photo)
        except Exception:
            pass

    def report_callback_exception(self, exc, val, tb) -> None:  # noqa: ANN001
        """Перехват ЛЮБОГО исключения в Tk-колбэке: логируем и продолжаем работу.

        Иначе в оконном .exe (где sys.stderr недоступен) необработанное исключение
        в колбэке может уронить приложение целиком. Здесь оно становится нефатальным.
        """
        import traceback

        text = "".join(traceback.format_exception(exc, val, tb))
        try:
            logger.error("Tk callback exception (не критично):\n%s", text)
        except Exception:
            pass
        try:
            if not self._closing and self.winfo_exists():
                self._set_status("Внутренняя ошибка интерфейса (см. лог) — приложение продолжит работу", _STATUS_WARN)
        except Exception:
            pass

    def _on_close(self) -> None:
        # Закрываемся чисто: глушим таймеры (иначе Tk ждёт отложенные after-колбэки
        # и процесс «висит»), затем best-effort выгружаем загруженные нами модели,
        # чтобы инстансы не копились в LM Studio после выхода.
        self._closing = True
        for attr in ("_model_poll_after_id", "_queue_poll_after_id", "_log_flush_after_id"):
            aid = getattr(self, attr, None)
            if aid is not None:
                try:
                    self.after_cancel(aid)
                except Exception:
                    pass
                setattr(self, attr, None)
        self._unload_models_on_close()
        self.destroy()

    def _unload_models_on_close(self) -> None:
        """Выгрузить модели, загруженные приложением (нативный LM Studio).

        Делаем в фоновом потоке с коротким join-таймаутом, чтобы недоступный сервер
        не подвесил выход. Управляется env NOCTURNE_UNLOAD_ON_CLOSE (1 по умолчанию).
        """
        if os.environ.get("NOCTURNE_UNLOAD_ON_CLOSE", "1").strip().lower() in ("0", "false", "no"):
            return
        try:
            import processor

            if not processor.app_loaded_models():
                return
            base_url = self._url_var.get().strip() or API_BASE
            api_key = self._api_key_var.get().strip() or API_KEY

            def _worker() -> None:
                try:
                    processor.unload_app_models(base_url, api_key)
                except Exception:
                    pass

            t = threading.Thread(target=_worker, daemon=True)
            t.start()
            t.join(timeout=4.0)
        except Exception:
            pass

    def _poll_loaded_model(self) -> None:
        if self._closing or not self.winfo_exists():
            return
        self._model_poll_active = True  # занять слот цикла (старт идёт в фоновом потоке)
        base_url = self._url_var.get().strip() or API_BASE
        api_key = self._api_key_var.get().strip() or API_KEY
        root = lmstudio_root_url(base_url)
        headers: dict[str, str] = {}
        if api_key:
            headers["Authorization"] = f"Bearer {api_key}"

        # Сетевой запрос — в фоновом потоке, чтобы не морозить UI (раньше блокировал
        # главный поток на timeout=5s каждые 4с).
        def _fetch() -> None:
            text = "недоступно"
            try:
                r = httpx.get(f"{root}/api/v1/models", headers=headers, timeout=5.0)
                if r.status_code < 400:
                    data = r.json()
                    loaded_names: list[str] = []
                    for m in data.get("models", []):
                        key = str(m.get("key") or m.get("id") or "")
                        loaded = m.get("loaded_instances") or []
                        if isinstance(loaded, list) and loaded:
                            loaded_names.append(key)
                    text = ", ".join(loaded_names) if loaded_names else "нет загруженных моделей"
            except Exception:
                text = "недоступно"
            if self._closing or not self.winfo_exists():
                return
            try:  # окно могло закрыться, пока ждали ответ — не шуметь трейсбэком
                self.after(0, lambda: self._apply_loaded_model_label(text))
            except (RuntimeError, tkinter.TclError):
                pass

        threading.Thread(target=_fetch, daemon=True).start()

    def _apply_loaded_model_label(self, text: str) -> None:
        if self._closing or not self.winfo_exists():
            return
        try:
            self._loaded_model_label.configure(text=f"Активная модель в LM Studio: {text}")
        except Exception:
            pass
        # На вкладке «Блокноты» панель анализа (и эта метка) скрыта — нет смысла
        # каждые 4с долбить сервер GET /api/v1/models. Ставим поллинг на паузу,
        # пока метка не видна; он возобновится при возврате на анализ-вкладку.
        if not getattr(self, "_analysis_panel", None) or not self._analysis_panel.winfo_ismapped():
            self._model_poll_active = False
            self._model_poll_after_id = None
            return
        self._model_poll_after_id = self.after(4000, self._poll_loaded_model)

    def _on_build_index(self) -> None:
        if self._running:
            self._set_status("Дождитесь завершения текущей задачи", _STATUS_WARN)
            return
        if self._folder_path is None and self._file_path is None:
            self._set_status("Выберите файл или папку для индекса", _STATUS_WARN)
            return
        emb = self._pick_embedding_model()
        if not emb:
            self._set_status("Не выбрана embedding-модель", _STATUS_WARN)
            return
        index_dir = Path(self._rag_index_dir_var.get().strip() or ".nocturne_index")
        base_url = self._url_var.get().strip() or API_BASE
        api_key = self._api_key_var.get().strip() or API_KEY
        context_budget = self._get_context_budget()
        reserve = self._get_response_reserve(context_budget)
        chunk_size = compute_dynamic_chunk_size(context_budget, SYSTEM_PROMPT_MAP, "index build", response_reserve=reserve)
        input_paths = [self._folder_path] if self._folder_path else [self._file_path]  # type: ignore[list-item]

        self._set_status("RAG: строю индекс…")
        self._append_log_line(f"[RAG] build_index dir={index_dir} embedding={emb}", "general")

        def do_build() -> None:
            try:
                stats = build_index(
                    input_paths=input_paths,  # type: ignore[arg-type]
                    index_dir=index_dir,
                    base_url=base_url,
                    api_key=api_key,
                    embedding_model=emb,
                    chunk_size_tokens=chunk_size,
                )
                self.after(0, lambda: self._set_status(
                    f"Индекс готов: chunks={stats.chunks_total}, files={stats.files_total}", _STATUS_OK
                ))
                self.after(0, lambda: self._append_log_line(
                    f"[RAG] index built dir={stats.index_dir} chunks={stats.chunks_total} files={stats.files_total}",
                    "summary",
                ))
            except Exception as exc:
                safe = sanitize_for_log(str(exc))
                self.after(0, lambda: self._set_status(f"Ошибка индекса: {safe}", _STATUS_ERR))

        threading.Thread(target=do_build, daemon=True).start()

    def _on_rag_ask(self) -> None:
        question = self._rag_question_text.get("1.0", "end-1c").strip()
        if not question:
            self._set_status("Введите вопрос для RAG", _STATUS_WARN)
            return
        emb = self._pick_embedding_model()
        if not emb:
            self._set_status("Не выбрана embedding-модель", _STATUS_WARN)
            return
        model = self._model_var.get().strip()
        if not model or model.startswith("("):
            self._set_status("Выберите LLM модель", _STATUS_WARN)
            return
        try:
            top_k = max(1, int(self._rag_top_k_var.get().strip() or "8"))
        except Exception:
            top_k = 8
        index_dir = Path(self._rag_index_dir_var.get().strip() or ".nocturne_index")
        base_url = self._url_var.get().strip() or API_BASE
        api_key = self._api_key_var.get().strip() or API_KEY

        self._set_status("RAG: ищу контекст и формирую ответ…")
        self._append_log_line(f"[RAG] query top_k={top_k} model={model} embedding={emb}", "general")

        def do_ask() -> None:
            try:
                hits = query_index(
                    question=question,
                    index_dir=index_dir,
                    base_url=base_url,
                    api_key=api_key,
                    embedding_model=emb,
                    top_k=top_k,
                )
                contexts = [h.text for h in hits]
                if not contexts:
                    self.after(0, lambda: self._rag_answer_text.delete("1.0", "end"))
                    self.after(0, lambda: self._rag_answer_text.insert("1.0", "Контексты не найдены в индексе."))
                    self.after(0, lambda: self._set_status("RAG: контексты не найдены", _STATUS_WARN))
                    return
                loop = asyncio.new_event_loop()
                asyncio.set_event_loop(loop)
                try:
                    answer = loop.run_until_complete(
                        answer_with_context(
                            question=question,
                            contexts=contexts,
                            base_url=base_url,
                            api_key=api_key,
                            model=model,
                            workers=max(1, int(self._workers_var.get() or "1")),
                            api_mode=self._api_mode_var.get().strip().lower(),
                        )
                    )
                finally:
                    loop.close()
                self.after(0, lambda: self._rag_answer_text.delete("1.0", "end"))
                self.after(0, lambda a=answer: self._rag_answer_text.insert("1.0", a))
                self.after(0, lambda: self._set_status("RAG: ответ готов", _STATUS_OK))
            except Exception as exc:
                safe = sanitize_for_log(str(exc))
                self.after(0, lambda: self._set_status(f"RAG ошибка: {safe}", _STATUS_ERR))

        threading.Thread(target=do_ask, daemon=True).start()

    def _on_recalc_selected_model_context(self) -> None:
        model = self._model_var.get().strip()
        if not model or model.startswith("("):
            self._set_status("Сначала выберите модель", _STATUS_WARN)
            return
        base_url = self._url_var.get().strip() or API_BASE
        api_key = self._api_key_var.get().strip() or API_KEY

        def recalc() -> None:
            self.after(0, lambda: self._set_status(f"Пересчитываю runtime-контекст: {model}…"))
            ctx, source, state = resolve_runtime_model_context(
                base_url=base_url,
                api_key=api_key,
                model=model,
                wait_for_loaded=True,
                poll_interval_seconds=1.0,
                stop_flag=lambda: self._closing,
            )
            if ctx:
                self._model_ctx[model] = ctx
                self._model_ctx_source[model] = source
                self.after(0, lambda: self._update_ctx_label(model))
                self.after(
                    0,
                    lambda: self._set_status(
                        f"Контекст {model}: {ctx:,} (source={source}, state={state})".replace(",", " "),
                        _STATUS_OK,
                    ),
                )
                self.after(
                    0,
                    lambda: self._append_log_line(
                        f"[RECALC] model={model} context={ctx} source={source} state={state}",
                        "general",
                    ),
                )
            else:
                self.after(
                    0,
                    lambda: self._set_status(
                        f"Не удалось пересчитать контекст {model} (state={state})",
                        _STATUS_WARN,
                    ),
                )

        threading.Thread(target=recalc, daemon=True).start()

    def _apply_runtime_state_to_widgets(self) -> None:
        state = self._runtime_state
        try:
            w = int(state.get("workers", 3))
        except Exception:
            w = 3
        self._workers_var.set(str(max(1, min(MAX_UI_WORKERS, w))))
        if str(state.get("base_url") or "").strip():
            self._url_var.set(str(state.get("base_url") or "").strip())
        self._composer_use_var.set(bool(state.get("composer_enabled", False)))
        self._api_mode_var.set("openai" if state.get("api_mode") == "openai" else "native")
        # Преселект провайдера под загруженные URL+режим (без перезаписи URL).
        try:
            preset = cp.get_preset(
                cp.detect_preset(self._url_var.get().strip(), self._api_mode_var.get().strip())
            )
            if preset is not None:
                self._provider_var.set(preset.label)
                self._provider_hint.configure(text=preset.hint)
        except Exception:
            pass
        self._low_vram_var.set(bool(state.get("low_vram_mode", True)))
        self._max_reduce_tokens_var.set(str(state.get("max_reduce_input_tokens", 24000)))
        self._max_chunk_tokens_var.set(str(state.get("max_chunk_tokens", 6000)))
        self._scout_var.set(bool(state.get("scout_mode", False)))
        self._scout_threshold_var.set(str(state.get("scout_threshold", 0.35)))
        sm = str(state.get("selected_scout_model") or "").strip()
        if sm:
            self._scout_model_var.set(sm)
        self._rag_index_dir_var.set(str(state.get("rag_index_dir") or ".nocturne_index"))
        self._rag_top_k_var.set(str(state.get("rag_top_k", 8)))
        em = str(state.get("selected_embedding_model") or "").strip()
        if em:
            self._embedding_model_var.set(em)
        mcc = state.get("model_context_cache")
        if isinstance(mcc, dict) and mcc:
            try:
                self._model_ctx.update({str(k): int(v) for k, v in mcc.items()})
            except (TypeError, ValueError):
                pass
        try:
            set_runtime_limits(
                max_reduce_input_tokens=int(self._max_reduce_tokens_var.get() or "24000"),
                max_chunk_tokens=int(self._max_chunk_tokens_var.get() or "6000"),
            )
        except Exception:
            pass
        self._on_composer_toggle()

    def _collect_runtime_state(self) -> dict[str, object]:
        try:
            workers_val = int(self._workers_var.get() or "3")
        except ValueError:
            workers_val = 3
        try:
            max_reduce_tokens = int(self._max_reduce_tokens_var.get() or "24000")
        except ValueError:
            max_reduce_tokens = 24000
        try:
            max_chunk_tokens = int(self._max_chunk_tokens_var.get() or "6000")
        except ValueError:
            max_chunk_tokens = 6000
        try:
            rag_top_k = int(self._rag_top_k_var.get() or "8")
        except ValueError:
            rag_top_k = 8
        return {
            "selected_model": self._model_var.get().strip(),
            "selected_vision_model": self._vision_model_var.get().strip(),
            "selected_composer_model": self._composer_model_var.get().strip(),
            "selected_embedding_model": self._embedding_model_var.get().strip(),
            "composer_enabled": bool(self._composer_use_var.get()),
            "workers": max(1, min(MAX_UI_WORKERS, workers_val)),
            "api_mode": self._api_mode_var.get().strip().lower(),
            "low_vram_mode": bool(self._low_vram_var.get()),
            "base_url": self._url_var.get().strip(),
            "max_reduce_input_tokens": max_reduce_tokens,
            "max_chunk_tokens": max_chunk_tokens,
            "scout_mode": bool(self._scout_var.get()),
            "scout_threshold": self._scout_threshold_var.get().strip() or "0.35",
            "selected_scout_model": self._scout_model_var.get().strip(),
            "model_context_cache": dict(self._model_ctx),
            "rag_index_dir": self._rag_index_dir_var.get().strip() or ".nocturne_index",
            "rag_top_k": rag_top_k,
        }

    def _apply_run_profile(self, name: str) -> None:
        from run_profiles import get_profile

        prof = get_profile(name)
        self._scout_var.set(bool(prof.get("scout_mode", False)))
        self._scout_threshold_var.set(str(prof.get("scout_threshold", 0.35)))
        self._workers_var.set(str(prof.get("workers", 3)))
        self._max_chunk_tokens_var.set(str(prof.get("max_chunk_tokens", 6000)))
        self._composer_use_var.set(bool(prof.get("composer_enabled", False)))
        self._on_composer_toggle()
        try:
            set_runtime_limits(max_chunk_tokens=int(prof.get("max_chunk_tokens", 6000)))
        except Exception:
            pass
        self._persist_runtime_state()
        self._update_preflight_label()
        self._set_status(f"Режим: {name}", _STATUS_OK)

    def _on_large_corpus_preset(self) -> None:
        """Пресет для анализа очень больших папок/архивов: scout + composer + меньшие чанки."""
        self._apply_run_profile("large_corpus")
        cm = self._model_var.get().strip()
        if cm and not cm.startswith("("):
            if not self._composer_model_var.get().strip() or self._composer_model_var.get().startswith("("):
                self._composer_model_var.set(cm)

    def _toggle_advanced_panel(self) -> None:
        pad = {"padx": 12}
        if self._advanced_visible.get():
            self._advanced_frame.pack(fill="x", pady=(0, 4), **pad)
        else:
            self._advanced_frame.pack_forget()

    def _apply_source_path(self, path: Path) -> None:
        if path.is_dir():
            self._folder_path = path
            self._file_path = None
            self._file_label.configure(text=f"[Папка]  {path}")
        else:
            self._file_path = path
            self._folder_path = None
            self._file_label.configure(text=str(path))
        self._hint_large_corpus(path)
        self._update_preflight_label()

    def _on_resume_job(self) -> None:
        from cache import get_job_state, list_resumable_jobs, load_last_job_pointer

        jobs = list_resumable_jobs(12)
        if not jobs:
            last = load_last_job_pointer()
            if last and last.get("job_id"):
                st = get_job_state(str(last["job_id"]))
                if st:
                    cached = 0
                    try:
                        from cache import count_cached_chunks

                        cached = count_cached_chunks(str(last["job_id"]))
                    except Exception:
                        pass
                    total = int(st.get("chunks_total") or 0)
                    if total > 0 and cached < total:
                        jobs = [{
                            "job_id": last["job_id"],
                            "query_preview": last.get("query_preview", ""),
                            "source_path": last.get("source_path", ""),
                            "chunks_total": total,
                            "cached": cached,
                            "status": st.get("status", "running"),
                        }]
        if not jobs:
            self._set_status("Нет незавершённых задач для продолжения", _STATUS_WARN)
            return

        dlg = ctk.CTkToplevel(self)
        dlg.title("Продолжить задачу")
        dlg.geometry("720x360")
        ctk.CTkLabel(
            dlg,
            text="Выберите прогон (MAP-кэш). Путь и запрос подставятся автоматически.",
            wraplength=680,
            justify="left",
        ).pack(anchor="w", padx=12, pady=(12, 8))

        def _pick(job: dict[str, object]) -> None:
            src = str(job.get("source_path") or "").strip()
            if not src or not Path(src).exists():
                self._set_status("Исходный путь недоступен — выберите тот же файл/папку вручную", _STATUS_WARN)
                dlg.destroy()
                return
            self._apply_source_path(Path(src))
            q = str(job.get("query_preview") or "").strip()
            if q:
                self._query_text.delete("1.0", "end")
                self._query_text.insert("1.0", q)
            cached = int(job.get("cached") or 0)
            total = int(job.get("chunks_total") or 0)
            self._append_log_line(
                f"[RESUME] job={str(job.get('job_id',''))[:16]}… MAP {cached}/{total}",
                "preflight",
            )
            dlg.destroy()
            self._on_start()

        scroll = ctk.CTkScrollableFrame(dlg, height=240)
        scroll.pack(fill="both", expand=True, padx=12, pady=(0, 12))
        for job in jobs:
            cached = int(job.get("cached") or 0)
            total = int(job.get("chunks_total") or 0)
            src = str(job.get("source_path") or "")
            preview = str(job.get("query_preview") or "")[:120]
            row = ctk.CTkFrame(scroll, fg_color=_md3.SURFACE_CONTAINER)
            row.pack(fill="x", pady=4)
            ctk.CTkLabel(
                row,
                text=f"{cached}/{total} чанков  |  {Path(src).name}\n{preview}",
                anchor="w",
                justify="left",
                wraplength=520,
            ).pack(side="left", fill="x", expand=True, padx=8, pady=8)
            ctk.CTkButton(
                row, text="Продолжить", width=100,
                command=lambda j=job: _pick(j),
            ).pack(side="right", padx=8, pady=8)

    def _on_dry_run(self) -> None:
        selected = self._file_path or self._folder_path
        if not selected:
            self._set_status("Выберите файл или папку для оценки", _STATUS_WARN)
            return
        query = self._query_text.get("1.0", "end").strip() or "(без запроса)"
        model = self._model_var.get().strip()
        ctx = self._model_ctx.get(model, CONTEXT_FALLBACK) if model and not model.startswith("(") else CONTEXT_FALLBACK
        reserve = self._get_response_reserve(ctx)
        try:
            chunk = compute_dynamic_chunk_size(ctx, SYSTEM_PROMPT_MAP, query, response_reserve=reserve)
        except Exception:
            chunk = 4500

        def work() -> None:
            try:
                from corpus_planner import format_plan_ru, plan_corpus

                plan = plan_corpus(
                    selected,
                    query,
                    chunk,
                    scout_mode=bool(self._scout_var.get()),
                    scout_threshold=float(self._scout_threshold_var.get() or "0.35"),
                )
                text = format_plan_ru(plan)
                self.after(0, lambda: self._preflight_label.configure(text=f"Оценка: {text}"))
                self.after(0, lambda: self._set_status("Dry-run готов (без вызовов LLM)", _STATUS_OK))
            except Exception as exc:
                safe = sanitize_for_log(str(exc))
                self.after(0, lambda: self._set_status(f"Оценка: {safe}", _STATUS_ERR))

        self._set_status("Считаю чанки и объём…")
        threading.Thread(target=work, daemon=True).start()

    def _on_run_history(self) -> None:
        from metrics import list_recent_runs

        runs = list_recent_runs(15)
        dlg = ctk.CTkToplevel(self)
        dlg.title("История прогонов")
        dlg.geometry("640x320")
        if not runs:
            ctk.CTkLabel(dlg, text="Пока нет сохранённых прогонов").pack(pady=20)
            return
        box = ctk.CTkTextbox(dlg, wrap="none")
        box.pack(fill="both", expand=True, padx=10, pady=10)
        for r in runs:
            line = (
                f"#{r['id']} job={r.get('job_id','')[:12]}… "
                f"ok={r.get('chunks_ok')}/{r.get('chunks_failed')} "
                f"scout_skip={r.get('scout_skipped')} "
                f"{r.get('duration_s') or 0:.0f}s | {r.get('query_preview','')}\n"
            )
            box.insert("end", line)
        box.configure(state="disabled")

    def _persist_runtime_state(self) -> None:
        if not self._runtime_state_ready:
            return
        try:
            try:
                set_runtime_limits(
                    max_reduce_input_tokens=int(self._max_reduce_tokens_var.get() or "24000"),
                    max_chunk_tokens=int(self._max_chunk_tokens_var.get() or "6000"),
                )
            except Exception:
                pass
            state = self._collect_runtime_state()
            self._runtime_state = dict(state)
            save_ui_runtime_state(state)
        except Exception as exc:
            logger.warning("Cannot persist runtime ui state: %s", sanitize_for_log(str(exc)))

    def _on_runtime_mode_changed(self) -> None:
        set_runtime_modes(
            api_mode=self._api_mode_var.get().strip().lower(),
            low_vram_mode=bool(self._low_vram_var.get()),
            dual_instance_mode=False,
        )
        self._persist_runtime_state()

    def _on_model_changed(self, *_args: Any) -> None:
        self._update_ctx_label(self._model_var.get())
        self._persist_runtime_state()

    def _on_composer_toggle(self) -> None:
        enabled = bool(self._composer_use_var.get())
        self._composer_menu.configure(state="normal" if enabled else "disabled")
        self._persist_runtime_state()

    def _bind_runtime_state_watchers(self) -> None:
        self._model_var.trace_add("write", self._on_model_changed)
        self._vision_model_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._composer_model_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._embedding_model_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._composer_use_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._workers_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._low_vram_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._url_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._max_reduce_tokens_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._max_chunk_tokens_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._rag_index_dir_var.trace_add("write", lambda *_: self._persist_runtime_state())
        self._rag_top_k_var.trace_add("write", lambda *_: self._persist_runtime_state())

    def _validate_lm_connection_fields(self) -> tuple[bool, str]:
        """Проверка полей подключения до сетевых запросов."""
        return validate_lmstudio_url(self._url_var.get())

    def _on_fetch_models(self) -> None:
        url = self._url_var.get().strip()
        key = self._api_key_var.get().strip()
        mode_state = set_runtime_modes(
            api_mode=self._api_mode_var.get().strip().lower(),
            low_vram_mode=bool(self._low_vram_var.get()),
            dual_instance_mode=False,
        )
        ok, err = self._validate_lm_connection_fields()
        if not ok:
            self._set_status(err, _STATUS_WARN)
            return

        def do_fetch() -> None:
            try:
                models, ctx, sources = fetch_models_info(url, key)
                by_kind = categorize_models(url, key)
                cat_tokens = summarize_model_tokens_by_category(url, key)
                self.after(0, lambda m=models, c=ctx, s=sources, k=by_kind: self._set_models(m, c, s, k))
                self.after(
                    0,
                    lambda t=cat_tokens, k=by_kind: self._append_log_line(
                        f"[MODELS tokens] llm={t.get('llm', 0)} embedding={t.get('embedding', 0)} "
                        f"vision={t.get('vision', 0)} tool={t.get('tool', 0)} "
                        f"reasoning={len(k.get('reasoning', []))}",
                        "general",
                    ),
                )
            except Exception as exc:
                safe = sanitize_for_log(str(exc))
                self.after(0, lambda s=safe: self._set_status(f"Ошибка: {s}", _STATUS_ERR))

        self._set_status(f"Загружаю список моделей… ({mode_state.get('api_mode', 'native')})")
        threading.Thread(target=do_fetch, daemon=True).start()

    def _set_models(
        self,
        models: list[str],
        ctx: dict[str, int],
        sources: dict[str, str],
        by_kind: dict[str, list[str]] | None = None,
    ) -> None:
        self._model_ctx = ctx
        self._model_ctx_source = sources
        if by_kind:
            self._models_by_kind = by_kind
        vals = models if models else ["(модели не найдены)"]
        chat_vals = self._models_by_kind.get("chat") or vals
        vision_vals = self._models_by_kind.get("vision") or chat_vals
        embed_vals = self._models_by_kind.get("embedding") or vals
        self._model_menu.configure(values=chat_vals if chat_vals else vals)
        self._vision_menu.configure(values=vision_vals if vision_vals else vals)
        self._composer_menu.configure(values=chat_vals if chat_vals else vals)
        self._embedding_menu.configure(values=embed_vals if embed_vals else vals)
        scout_vals = ["(как MAP-модель)"] + (models if models else [])
        self._scout_menu.configure(values=scout_vals)
        if models:
            persisted_main = str(self._runtime_state.get("selected_model") or "").strip()
            persisted_vision = str(self._runtime_state.get("selected_vision_model") or "").strip()
            persisted_composer = str(self._runtime_state.get("selected_composer_model") or "").strip()
            persisted_embedding = str(self._runtime_state.get("selected_embedding_model") or "").strip()
            dm = self._cfg_default_model
            if persisted_main and persisted_main in models:
                chat = persisted_main
            elif dm and dm in models:
                chat = dm
            else:
                chat = next((m for m in models if "embed" not in m.lower()), models[0])
            vision_sel = persisted_vision if persisted_vision in models else chat
            composer_sel = persisted_composer if persisted_composer in models else chat
            embedding_sel = (
                persisted_embedding
                if persisted_embedding in models
                else next((m for m in models if "embed" in m.lower()), chat)
            )
            self._model_var.set(chat)
            self._vision_model_var.set(vision_sel)
            self._composer_model_var.set(composer_sel)
            self._embedding_model_var.set(embedding_sel)
            self._update_ctx_label(chat)
        n_reason = len(self._models_by_kind.get("reasoning") or [])
        status = f"Загружено моделей: {len(models)}"
        if n_reason:
            status += f" (reasoning: {n_reason})"
        self._set_status(status)
        self._persist_runtime_state()
        if not self._model_poll_active:
            self._poll_loaded_model()

    def _update_ctx_label(self, model: str) -> None:
        ctx = self._model_ctx.get(model)
        source = self._model_ctx_source.get(model, "fallback")
        if ctx:
            source_hint = (
                "runtime loaded" if source == "runtime_loaded"
                else "metadata (not loaded)" if source == "metadata_not_loaded"
                else "из /models" if source == "metadata"
                else "через probe" if source == "probe"
                else f"fallback {CONTEXT_FALLBACK:,}".replace(",", " ")
            )
            self._ctx_label.configure(
                text=f"Контекст модели: {ctx:,} токенов ({source_hint})".replace(",", " "),
                text_color=_STATUS_OK,
            )
        else:
            self._ctx_label.configure(
                text=(
                    "Контекст модели: не определён, используется fallback "
                    f"{CONTEXT_FALLBACK:,}".replace(",", " ")
                ),
                text_color=_STATUS_WARN,
            )

    def _get_context_budget(self) -> int:
        """Бюджет контекста только из LM Studio metadata."""
        model = self._model_var.get().strip()
        return max(500, self._model_ctx.get(model, CONTEXT_FALLBACK))

    def _get_response_reserve(self, context_budget: int) -> int:
        # Автоматический резерв без ручной настройки: ~20% контекста, но в разумных рамках.
        reserve = max(1024, min(4096, int(context_budget * 0.2)))
        return max(256, min(reserve, max(256, context_budget - 500)))

    def _on_provider_change(self, label: str) -> None:
        preset = cp.preset_by_label(label)
        if preset is None:
            return
        # На старте (авто-детект) только показываем подсказку — НЕ перетираем
        # сохранённый пользователем URL/режим.
        if not getattr(self, "_runtime_state_ready", False):
            self._provider_hint.configure(text=preset.hint)
            return
        if preset.autofills_url:
            self._url_var.set(preset.base_url)
        self._api_mode_var.set(preset.api_mode)
        self._provider_hint.configure(text=preset.hint)
        # Применить и сохранить, как при ручной смене режима.
        self._on_runtime_mode_changed()
        self._persist_runtime_state()
        self._append_log_line(
            f"[PRESET] {preset.label}: api_mode={preset.api_mode}, "
            f"url={preset.base_url if preset.autofills_url else '(вручную)'}",
            "preflight",
        )
        self._set_status(f"Провайдер: {preset.label}", _STATUS_OK)

    def _on_download_model_dialog(self) -> None:
        """Скачать модель в LM Studio (REST v1 /api/v1/models/download + status)."""
        base_url = self._url_var.get().strip() or API_BASE
        api_key = self._api_key_var.get().strip() or API_KEY
        root = lmstudio_root_url(base_url)

        dlg = ctk.CTkToplevel(self)
        dlg.title("Скачать модель (LM Studio)")
        dlg.geometry("460x230")
        ctk.CTkLabel(
            dlg, text="Загрузка модели в LM Studio",
            font=ctk.CTkFont(size=14, weight="bold"),
        ).pack(pady=(12, 4))
        ctk.CTkLabel(
            dlg,
            text="Идентификатор модели (например, repo из каталога LM Studio / HF).",
            text_color=_md3.ON_SURFACE_VARIANT, font=ctk.CTkFont(size=11), wraplength=420, justify="left",
        ).pack(pady=(0, 6))
        model_var = ctk.StringVar(value="")
        ctk.CTkEntry(dlg, textvariable=model_var, width=420).pack(pady=4)
        status = ctk.CTkLabel(dlg, text="", text_color=_md3.ON_SURFACE_VARIANT, wraplength=420, justify="left")
        status.pack(pady=6)

        def _set(text: str, color: str = "gray") -> None:
            if dlg.winfo_exists():
                status.configure(text=text, text_color=color)

        def _poll(job_id: str) -> None:
            if self._closing or not dlg.winfo_exists():
                return  # диалог закрыт — прекращаем опрос (загрузка продолжится на сервере)
            st = lmsapi.get_model_download_status(root, api_key, job_id)
            if "error" in st:
                self.after(0, lambda: _set(f"Статус недоступен: {st['error']}", _STATUS_WARN))
                return
            state = str(st.get("status") or st.get("state") or "").lower()
            prog = st.get("progress")
            pct = f" {float(prog) * 100:.0f}%" if isinstance(prog, (int, float)) else ""
            if state in ("completed", "done", "success", "finished"):
                self.after(0, lambda: _set("Готово ✓ — обновите список моделей.", _STATUS_OK))
                return
            if state in ("failed", "error", "cancelled", "canceled"):
                self.after(0, lambda: _set(f"Не удалось: {sanitize_for_log(str(st))[:200]}", _STATUS_ERR))
                return
            self.after(0, lambda: _set(f"Загрузка…{pct} ({state or 'в процессе'})"))
            if not self._closing and self.winfo_exists() and dlg.winfo_exists():
                self.after(2000, lambda: threading.Thread(
                    target=_poll, args=(job_id,), daemon=True).start())

        def _start() -> None:
            model = model_var.get().strip()
            if not model:
                _set("Укажите идентификатор модели.", _STATUS_WARN)
                return
            _set("Запуск загрузки…")

            def _worker() -> None:
                resp = lmsapi.start_model_download(root, api_key, model)
                if "error" in resp:
                    self.after(0, lambda: _set(f"Ошибка: {resp['error']}", _STATUS_ERR))
                    return
                job_id = str(resp.get("job_id") or resp.get("id") or "").strip()
                if not job_id:
                    self.after(0, lambda: _set(
                        "Загрузка запущена (job_id не вернулся). "
                        "Проверьте прогресс в LM Studio.", _STATUS_OK))
                    return
                threading.Thread(target=_poll, args=(job_id,), daemon=True).start()

            threading.Thread(target=_worker, daemon=True).start()

        ctk.CTkButton(dlg, text="Скачать", command=_start).pack(pady=10)

    def _on_test_connection(self) -> None:
        set_runtime_modes(
            api_mode=self._api_mode_var.get().strip().lower(),
            low_vram_mode=bool(self._low_vram_var.get()),
            dual_instance_mode=False,
        )
        ok, err = self._validate_lm_connection_fields()
        if not ok:
            self._set_status(err, _STATUS_WARN)
            return
        base_url = self._url_var.get().strip() or API_BASE
        api_key  = self._api_key_var.get().strip() or API_KEY

        def do_test() -> None:
            chat = self._model_var.get().strip()
            emb = self._pick_embedding_model() or None
            full = bool(chat and not chat.startswith("("))
            ok, msg = test_lmstudio_connection(
                base_url,
                api_key,
                embedding_model=emb,
                full_smoke=full,
                chat_model=chat if full else None,
            )
            msg = sanitize_for_log(msg)
            color = _STATUS_OK if ok else _STATUS_ERR
            self.after(0, lambda: self._status_label.configure(
                text=f"LM Studio: {msg}", text_color=color,
            ))

        self._set_status("Проверка LM Studio…")
        threading.Thread(target=do_test, daemon=True).start()

    def _on_test_vision(self) -> None:
        set_runtime_modes(
            api_mode=self._api_mode_var.get().strip().lower(),
            low_vram_mode=bool(self._low_vram_var.get()),
            dual_instance_mode=False,
        )
        ok, err = self._validate_lm_connection_fields()
        if not ok:
            self._set_status(err, _STATUS_WARN)
            return
        base_url = self._url_var.get().strip() or API_BASE
        api_key = self._api_key_var.get().strip() or API_KEY
        vm = self._vision_model_var.get().strip() or self._model_var.get().strip()
        if not vm or vm.startswith("("):
            self._set_status("Выберите Vision-модель или основную LLM", _STATUS_WARN)
            return

        def do_test() -> None:
            ok, msg = check_vision_capability(base_url, api_key, vm)
            msg = sanitize_for_log(msg)
            color = _STATUS_OK if ok else _STATUS_ERR
            self.after(0, lambda: self._append_log_line(f"[VISION check] {msg}", "vision_map"))
            self.after(0, lambda: self._status_label.configure(text=f"Vision: {msg}", text_color=color))

        self._set_status(f"Проверка Vision: {vm}…")
        threading.Thread(target=do_test, daemon=True).start()

    def _on_select_file(self) -> None:  # noqa: D401 - updates preflight
        path = ctk.filedialog.askopenfilename(filetypes=FILE_TYPES)
        if path:
            self._file_path = Path(path)
            self._folder_path = None
            self._file_label.configure(text=str(self._file_path))
            self._hint_large_corpus(self._file_path)
            self._update_preflight_label()

    def _on_select_folder(self) -> None:
        path = ctk.filedialog.askdirectory()
        if path:
            self._folder_path = Path(path)
            self._file_path = None
            self._file_label.configure(text=f"[Папка]  {self._folder_path}")
            self._hint_large_corpus(self._folder_path)
            self._update_preflight_label()

    def _hint_large_corpus(self, selected: Path) -> None:
        from large_corpus_io import is_large_corpus_input

        ok, reason = is_large_corpus_input(selected)
        if ok:
            self._set_status(
                f"Большой корпус ({reason}): при старте включим scout и пресет 1M+ автоматически",
                _md3.PRIMARY,
            )

    def _auto_apply_large_corpus_if_needed(self, selected: Path) -> str | None:
        """Применить пресет large_corpus без ручной донастройки."""
        from large_corpus_io import is_large_corpus_input, large_corpus_profile_kwargs

        ok, reason = is_large_corpus_input(selected)
        if not ok:
            return None
        prof = large_corpus_profile_kwargs()
        self._scout_var.set(bool(prof.get("scout_mode", True)))
        self._scout_threshold_var.set(str(prof.get("scout_threshold", 0.35)))
        self._workers_var.set(str(prof.get("workers", 4)))
        self._max_chunk_tokens_var.set(str(prof.get("max_chunk_tokens", 4500)))
        self._composer_use_var.set(bool(prof.get("composer_enabled", True)))
        self._on_composer_toggle()
        cm = self._model_var.get().strip()
        if cm and not cm.startswith("("):
            if not self._composer_model_var.get().strip() or self._composer_model_var.get().startswith("("):
                self._composer_model_var.set(cm)
        try:
            set_runtime_limits(max_chunk_tokens=int(prof.get("max_chunk_tokens", 4500)))
        except Exception:
            pass
        self._persist_runtime_state()
        return f"large_corpus auto ({reason}): scout, chunk={prof.get('max_chunk_tokens')}, workers={prof.get('workers')}"

    def _on_stop(self) -> None:
        self._stop_requested = True
        if self._active_job_id:
            try:
                from cache import mark_job_paused

                mark_job_paused(self._active_job_id)
            except Exception:
                pass
        self._set_status("Остановка после текущего чанка… (можно «Продолжить»)")

    def _on_start(self) -> None:
        if self._running:
            return

        selected = self._file_path or self._folder_path
        if not selected or not selected.exists():
            self._set_status("Выберите файл или папку")
            return

        query = self._query_text.get("1.0", "end").strip()
        if not query:
            self._set_status("Введите запрос")
            return

        ok_lm, err_lm = self._validate_lm_connection_fields()
        if not ok_lm:
            self._set_status(err_lm, _STATUS_WARN)
            return

        model = self._model_var.get().strip()
        if not model or model.startswith("("):
            self._set_status("Выберите модель (нажмите «Обновить модели»)")
            return

        auto_line = self._auto_apply_large_corpus_if_needed(selected)
        if auto_line:
            self._append_log_line(f"[AUTO] {auto_line}", "preflight")

        context_budget = self._get_context_budget()
        response_reserve = self._get_response_reserve(context_budget)
        try:
            workers = max(1, min(MAX_UI_WORKERS, int(self._workers_var.get())))
        except ValueError:
            workers = 3

        base_url = self._url_var.get().strip() or API_BASE
        api_key  = self._api_key_var.get().strip() or API_KEY

        # Reset state
        self._last_result_text = ""
        self._last_result_df   = None
        self._stop_requested   = False
        # job_id вычисляется внутри обработки (учитывает chunk_size/model/composer)
        # и приходит в GUI через MSG_JOB_ID. Здесь не предугадываем, иначе «Стоп»
        # до прихода MSG_JOB_ID пометил бы несуществующий job.
        self._active_job_id = None
        self._result_text.delete("1.0", "end")
        self._progress_bar.set(0)
        self._running = True
        self._start_btn.configure(state="disabled")
        self._stop_btn.configure(state="normal")
        # Hide meta-prompt preview from previous run
        self._meta_prompt_label.configure(text="")
        self._meta_prompt_label.pack_forget()
        self._set_status(
            (
                f"Старт: контекст={context_budget:,}, резерв={response_reserve:,}, "
                f"воркеров={workers}, chunk≈контекст-system-query-резерв"
            ).replace(",", " ")
        )
        vm = self._vision_model_var.get().strip() or None
        cm = (
            self._composer_model_var.get().strip()
            if bool(self._composer_use_var.get())
            else None
        )
        api_mode = self._api_mode_var.get().strip().lower()
        low_vram_mode = bool(self._low_vram_var.get())
        dual_instance_mode = False
        scout_mode = bool(self._scout_var.get())
        try:
            scout_threshold = float(self._scout_threshold_var.get().strip() or "0.35")
        except ValueError:
            scout_threshold = 0.35
        scout_threshold = max(0.0, min(1.0, scout_threshold))
        scout_m = self._scout_model_var.get().strip()
        if scout_m.startswith("(") or not scout_m:
            scout_m = None
        set_runtime_modes(
            api_mode=api_mode,
            low_vram_mode=low_vram_mode,
            dual_instance_mode=dual_instance_mode,
        )
        try:
            set_runtime_limits(
                max_reduce_input_tokens=int(self._max_reduce_tokens_var.get() or "24000"),
                max_chunk_tokens=int(self._max_chunk_tokens_var.get() or "6000"),
            )
        except Exception:
            pass
        self._persist_runtime_state()
        self._update_preflight_label()
        try:
            mct = int(self._max_chunk_tokens_var.get() or "6000")
        except ValueError:
            mct = 6000
        try:
            mrt = int(self._max_reduce_tokens_var.get() or "24000")
        except ValueError:
            mrt = 24000

        # Единая неизменяемая конфигурация прогона — носитель настроек GUI
        # для processing-слоя (после preflight уточняется через replace()).
        rc = RunConfig.from_gui(
            base_url=base_url,
            api_key=api_key,
            chat_model=model,
            vision_model=vm,
            composer_model=cm,
            scout_model=scout_m,
            embedding_model=self._embedding_model_var.get().strip(),
            api_mode=api_mode,
            low_vram=low_vram_mode,
            workers=workers,
            context_budget=context_budget,
            response_reserve=response_reserve,
            max_chunk_tokens=mct,
            max_reduce_input_tokens=mrt,
            scout_mode=scout_mode,
            scout_threshold=scout_threshold,
        )
        self._append_log_line(
            f"[RUN_CONFIG] scout={rc.scout_mode} workers={rc.workers} chunk_max={rc.max_chunk_tokens}",
            "preflight",
        )
        self._append_log_line(
            (
                f"[START] model={model} vision={vm or model} composer={cm or model} "
                f"context={context_budget} reserve={response_reserve} "
                f"workers={workers} source={self._model_ctx_source.get(model, 'unknown')} "
                f"api_mode={api_mode} low_vram={low_vram_mode} dual_instance={dual_instance_mode} "
                f"scout={scout_mode} scout_threshold={scout_threshold}"
            ),
            "general",
        )

        def run() -> None:
            try:
                # Runtime preflight: ждём фактическую загрузку выбранной модели и
                # берём реальный контекст из LM Studio (loaded_context_length).
                self._queue.put({
                    "type": MSG_PROGRESS,
                    "current": 0,
                    "total": 1,
                    "phase": "preflight",
                    "message": f"Проверяю runtime контекст модели: {model}",
                })
                runtime_ctx, runtime_source, runtime_state = resolve_runtime_model_context(
                    base_url=base_url,
                    api_key=api_key,
                    model=model,
                    wait_for_loaded=True,
                    poll_interval_seconds=1.0,
                    stop_flag=lambda: self._stop_requested or self._closing,
                )
                effective_context = runtime_ctx or context_budget
                effective_reserve = self._get_response_reserve(effective_context)
                # Уточняем конфиг реальным runtime-контекстом из preflight.
                from dataclasses import replace as _dc_replace

                rc_effective = _dc_replace(
                    rc,
                    context_budget=effective_context,
                    response_reserve=effective_reserve,
                )
                self._queue.put({
                    "type": MSG_PROGRESS,
                    "current": 0,
                    "total": 1,
                    "phase": "preflight",
                    "message": (
                        f"Контекст модели: {effective_context:,}, резерв: {effective_reserve:,} "
                        f"(source={runtime_source}, state={runtime_state or 'unknown'})"
                    ).replace(",", " "),
                })
                self._queue.put({
                    "type": MSG_TRACE,
                    "line": (
                        f"[PREFLIGHT] model={model} context={effective_context} reserve={effective_reserve} "
                        f"source={runtime_source} state={runtime_state or 'unknown'}"
                    ),
                })
                _run_processing(
                    file_path=self._file_path,
                    folder_path=self._folder_path,
                    query=query,
                    rc=rc_effective,
                    out_queue=self._queue,
                    stop_flag=lambda: self._stop_requested,
                )
            except Exception as exc:
                logger.exception("Worker thread error: %s", sanitize_for_log(str(exc)))
                self._queue.put({"type": MSG_ERROR, "message": sanitize_for_log(str(exc))})
            self._queue.put({"type": "done"})

        threading.Thread(target=run, daemon=True).start()
        self._poll_queue()

    # ------------------------------------------------------------------ #
    #  Queue polling
    # ------------------------------------------------------------------ #

    def _poll_queue(self) -> None:
        try:
            while True:
                msg = self._queue.get_nowait()
                t = msg.get("type")

                if t == MSG_PROGRESS:
                    cur   = msg.get("current", 0)
                    total = max(1, msg.get("total", 1))
                    phase = msg.get("phase", "")
                    cache = msg.get("from_cache", 0)
                    active = msg.get("active", 0)
                    in_flight = msg.get("in_flight", 0)
                    retrying = msg.get("retrying", 0)
                    effective_workers = msg.get("effective_workers", 0)
                    chunk_idx = msg.get("chunk_idx", 0)
                    file_name = msg.get("file", "")
                    preview = msg.get("preview", "")
                    route = msg.get("route", "single")
                    instance_id = msg.get("instance_id", "")
                    if phase != "map_started":
                        self._progress_bar.set(cur / total)
                    if phase == "extract":
                        suffix = f" | файл: {file_name}" if file_name else ""
                        self._set_status(f"Читаю файлы: {cur} / {total}…{suffix}")
                        self._append_log_line(f"[EXTRACT] {cur}/{total}{suffix}", "extract")
                    elif phase == "preflight":
                        self._set_status(msg.get("message", "Проверяю модель и контекст…"))
                    elif phase == "map_started":
                        suffix = f" | файл: {file_name}" if file_name else ""
                        self._set_status(
                            f"MAP: запущено {cur}/{total}, active={active}, in_flight={in_flight}, backoff={retrying}, eff={effective_workers}{suffix}"
                        )
                        if preview and chunk_idx:
                            self._append_log_line(
                                f"[MAP start #{chunk_idx}] route={route} instance={instance_id or '-'} active={active} in_flight={in_flight} retrying={retrying} eff={effective_workers} file={file_name or 'n/a'} | {preview}",
                                "map",
                            )
                    elif phase == "scout":
                        scout_total = msg.get("scout_total", total)
                        self._set_status(f"Scout: оценка релевантности 0/{scout_total}…")
                        self._append_log_line(
                            f"[SCOUT] starting relevance pass for {scout_total} chunks",
                            "map",
                        )
                    elif phase == "scout_done":
                        deep = msg.get("scout_deep", cur)
                        skipped = msg.get("scout_skipped", 0)
                        thr = msg.get("scout_threshold", 0.35)
                        self._set_status(
                            f"Scout готов: deep MAP {deep}, пропущено {skipped} (порог {thr})",
                            _STATUS_OK,
                        )
                        self._append_log_line(
                            f"[SCOUT done] deep_map={deep} skipped={skipped} threshold={thr}",
                            "map",
                        )
                    elif phase == "map_resume":
                        self._set_status(
                            f"Продолжение: {cache} чанков уже в кэше ({cur}/{total})"
                        )
                        self._append_log_line(
                            f"[RESUME cache] loaded={cache} total={total}",
                            "preflight",
                        )
                    elif phase == "map_batch":
                        phase_name = msg.get("phase_name", "")
                        batch_done = msg.get("batch_done", 0)
                        batch_total = msg.get("batch_total", 0)
                        self._set_status(
                            f"MAP {phase_name}: пакет {batch_done}/{batch_total} "
                            f"(всего {cur}/{total}, кэш {cache})"
                        )
                    elif phase == "map":
                        extra = f" ({cache} из кэша)" if cache else ""
                        suffix = f" | файл: {file_name}" if file_name else ""
                        self._set_status(
                            f"MAP: {cur}/{total}, active={active}, in_flight={in_flight}, backoff={retrying}, eff={effective_workers}{extra}{suffix}"
                        )
                        if chunk_idx:
                            self._append_log_line(
                                f"[MAP done #{chunk_idx}] route={route} instance={instance_id or '-'} done={cur}/{total} active={active} in_flight={in_flight} retrying={retrying} eff={effective_workers} cache={cache} file={file_name or 'n/a'}",
                                "map",
                            )
                    elif phase == "model_switch":
                        stage = msg.get("stage", "")
                        from_model = msg.get("from_model", "") or "none"
                        to_model = msg.get("to_model", "")
                        self._set_status(
                            f"Переключаю модель ({stage}): {from_model} -> {to_model}…"
                        )
                        self._append_log_line(
                            f"[MODEL switch] stage={stage} from={from_model} to={to_model}",
                            "preflight",
                        )
                    elif phase == "model_phase":
                        phase_name = msg.get("phase_name", "")
                        phase_model = msg.get("model", "")
                        chunks_count = msg.get("chunks_count", 0)
                        self._set_status(
                            f"Фаза {phase_name}: модель={phase_model}, чанков={chunks_count}"
                        )
                        self._append_log_line(
                            f"[MODEL phase] phase={phase_name} model={phase_model} chunks={chunks_count} active={active} in_flight={in_flight} retrying={retrying} eff={effective_workers}",
                            "preflight",
                        )
                    elif phase == "instance_pool":
                        model_name = msg.get("model", "")
                        ids = msg.get("instance_ids") or []
                        self._append_log_line(
                            f"[INSTANCE pool] model={model_name} count={msg.get('instances_loaded', 0)} ids={ids}",
                            "preflight",
                        )
                    elif phase == "vision_map":
                        vm = msg.get("vision_model", "")
                        self._append_log_line(
                            f"[VISION MAP #{chunk_idx}] model={vm} file={file_name or 'n/a'}",
                            "vision_map",
                        )
                    elif phase == "retry_scheduled":
                        attempt = msg.get("attempt", 0)
                        max_attempts = msg.get("max_attempts", 0)
                        error_kind = msg.get("error_kind", "unknown")
                        retry_delay = msg.get("retry_delay", 0.0)
                        self._append_log_line(
                            f"[RETRY] chunk={chunk_idx} route={route} instance={instance_id or '-'} {attempt}/{max_attempts} kind={error_kind} delay={retry_delay:.1f}s file={file_name or 'n/a'} active={active} in_flight={in_flight} retrying={retrying} eff={effective_workers} backoff_freed={msg.get('backoff_slot_freed', False)}",
                            "retry",
                        )
                    elif phase == "map_failed":
                        error_kind = msg.get("error_kind", "unknown")
                        err = str(msg.get("error", ""))[:220]
                        self._append_log_line(
                            f"[MAP failed #{chunk_idx}] route={route} instance={instance_id or '-'} kind={error_kind} classifier={msg.get('error_classifier', 'n/a')} file={file_name or 'n/a'} err={err}",
                            "error",
                        )
                    elif phase == "circuit_breaker":
                        pause_s = msg.get("pause_seconds", 0.0)
                        self._append_log_line(
                            f"[CIRCUIT BREAKER] pausing requests for {pause_s:.1f}s",
                            "retry",
                        )
                    elif phase == "summary":
                        retries = msg.get("retries", 0)
                        failed = msg.get("failed", 0)
                        ok = msg.get("ok", 0)
                        elapsed_s = float(msg.get("elapsed_s", 0.0) or 0.0)
                        cpm = float(msg.get("chunks_per_min", 0.0) or 0.0)
                        total_ch = int(msg.get("text_chunks", total) or total)
                        deep_ch = int(msg.get("text_map_chunks", ok) or ok)
                        if cpm > 0 and deep_ch > ok:
                            remain = max(0, deep_ch - cur)
                            eta_min = remain / cpm if cpm else 0
                            self._set_status(
                                f"MAP {cur}/{total_ch}: ~{eta_min:.0f} мин осталось "
                                f"({cpm:.1f} ch/min)",
                            )
                        wk = msg.get("workers", "")
                        mm = msg.get("map_model", "")
                        vm = msg.get("vision_model", "")
                        rm = msg.get("reduce_model", "")
                        self._append_log_line(
                            f"[SUMMARY] ok={ok} failed={failed} retries={retries} elapsed={elapsed_s:.1f}s "
                            f"throughput={cpm:.1f} chunks/min workers={wk} map={mm} vision={vm} reduce={rm} "
                            f"text_chunks={msg.get('text_chunks', 0)} text_map={msg.get('text_map_chunks', 0)} "
                            f"scout_skipped={msg.get('scout_skipped', 0)} scout_mode={msg.get('scout_mode', False)} "
                            f"vision_chunks={msg.get('vision_chunks', 0)} "
                            f"low_vram_sequential={msg.get('low_vram_sequential', True)} "
                            f"api_mode={msg.get('api_mode', 'native')} instances={msg.get('instances_loaded', 1)} "
                            f"dual_instance_active={msg.get('dual_instance_active', False)}",
                            "summary",
                        )
                    elif phase == "map_metrics":
                        rc = msg.get("relevant_chunks", 0)
                        fc = msg.get("findings_count", 0)
                        ev = msg.get("evidence_refs_count", 0)
                        self._append_log_line(
                            f"[MAP metrics] relevant_chunks={rc} findings={fc} evidence_refs={ev}",
                            "map_metrics",
                        )
                        self._set_status(
                            f"MAP: релевантных чанков {rc}, находок {fc}, evidence {ev}"
                        )
                    elif phase == "map_relevant":
                        rf = msg.get("relevant_files", cur)
                        self._append_log_line(
                            f"[MAP relevant] {rf} файловых групп с находками из {total} чанков",
                            "map_metrics",
                        )
                        self._set_status(f"MAP завершён: {rf} релевантных файловых групп")
                    elif phase == "section_reduce":
                        sf = msg.get("section_file", "")
                        self._set_status(f"REDUCE секции {cur}/{total}: {sf[:60]}")
                        self._append_log_line(
                            f"[SECTION REDUCE] {cur}/{total} file={sf}", "reduce"
                        )
                    elif phase == "synthesize":
                        self._set_status("Синтез финального отчёта…", _STATUS_OK)
                        self._append_log_line("[SYNTHESIZE] финальный синтез секций", "reduce")
                    elif phase == "reduce_merge":
                        ml = msg.get("merge_level", 0)
                        self._append_log_line(f"[REDUCE merge] {cur}/{total} level={ml}", "reduce_merge")
                    elif phase == "reduce_refine":
                        self._append_log_line("[REDUCE] второй проход refine (полнота отчёта)", "reduce")
                        self._set_status("REDUCE: refine — расширение отчёта…")
                    elif phase == "quality_metrics":
                        cc = msg.get("covered_chunks", 0)
                        ec = msg.get("evidence_count", 0)
                        sc = msg.get("final_sections_count", 0)
                        ru = msg.get("refine_used", False)
                        vw = msg.get("validation_warnings") or []
                        self._append_log_line(
                            f"[QUALITY] covered_chunks={cc} evidence_count={ec} "
                            f"sections={sc} refine={ru} warnings={vw}",
                            "quality_metrics",
                        )
                    elif phase == "meta_plan":
                        self._set_status("Composer: генерирую оптимальный промт для анализа…")
                        self._append_log_line("[META PLAN] composer generating analysis directive", "preflight")
                    elif phase == "meta_plan_done":
                        preview = str(msg.get("preview", ""))[:200]
                        self._append_log_line(f"[META PLAN done] directive={preview}…", "preflight")
                        self._set_status("Оптимальный промт готов, запускаю MAP…", _STATUS_OK)
                        # Show meta-prompt preview label
                        short = preview[:120] + ("…" if len(preview) >= 120 else "")
                        self._meta_prompt_label.configure(
                            text=f"Composer-промт: {short}"
                        )
                        self._meta_prompt_label.pack(anchor="w", pady=(0, 4))
                    elif phase in ("index_extract", "index_embed"):
                        label = "Индексирую файлы" if phase == "index_extract" else "Генерирую эмбеддинги"
                        self._set_status(f"RAG: {label}: {cur}/{total}…")
                        self._append_log_line(f"[RAG {phase}] {cur}/{total}", "general")
                    elif phase == "aggregate":
                        preview = str(msg.get("preview", ""))[:200]
                        self._set_status("Агрегация (детерминированная)…", _STATUS_OK)
                        self._append_log_line(f"[AGGREGATE] {preview}", "map_metrics")
                    elif phase == "run_diff":
                        preview = str(msg.get("preview", ""))[:200]
                        self._set_status("Сравнение с прошлым прогоном…", _STATUS_OK)
                        self._append_log_line(f"[RUN DIFF] {preview}", "map_metrics")
                    elif phase == "reduce":
                        self._set_status(f"REDUCE: группа {cur} / {total}…")
                        self._append_log_line(f"[REDUCE] {cur}/{total}", "reduce")
                    elif phase == "batch":
                        self._set_status(f"Батч: {cur} / {total}…")
                        self._append_log_line(f"[BATCH] {cur}/{total}", "batch")
                    elif phase == "stopped":
                        self._set_status(
                            f"Остановлено: {cur}/{total} чанков в кэше — "
                            "нажмите «Продолжить» для возобновления",
                            _STATUS_WARN,
                        )
                        self._append_log_line(
                            f"[STOPPED] map {cur}/{total} cached", "preflight",
                        )

                elif t == MSG_RESULT:
                    text = msg.get("text", "")
                    self._last_result_text = text
                    self._result_type = "text"
                    self._result_text.delete("1.0", "end")
                    self._result_text.insert("1.0", text)
                    self._progress_bar.set(1.0)
                    self._set_status("Готово ✓", _STATUS_OK)

                elif t == MSG_RESULT_DF:
                    df    = msg.get("df")
                    saved = msg.get("saved_path", "")
                    if isinstance(df, pd.DataFrame) and not df.empty:
                        self._last_result_df   = df
                        self._last_result_text = df.to_string()
                        self._result_type = "table"
                        self._result_text.delete("1.0", "end")
                        if saved:
                            self._result_text.insert("1.0", f"Сохранено: {saved}\n\n")
                        self._result_text.insert("end", self._last_result_text)
                    self._progress_bar.set(1.0)
                    self._set_status(
                        f"Готово ✓  →  {saved}" if saved else "Готово ✓",
                        _STATUS_OK,
                    )

                elif t == MSG_ERROR:
                    from errors import classify_exception

                    raw = str(msg.get("message", ""))
                    err = sanitize_for_log(raw)
                    hint = ""
                    try:
                        hint = classify_exception(Exception(raw)).user_hint()
                    except Exception:
                        pass
                    self._result_text.delete("1.0", "end")
                    self._result_text.insert("1.0", f"ОШИБКА:\n{err}\n\n{hint}")
                    self._set_status(f"Ошибка: {err[:100]}", _STATUS_ERR)
                    self._append_log_line(f"[ERROR] {err}", "error")

                elif t == MSG_JOB_ID:
                    jid = msg.get("job_id")
                    if isinstance(jid, str) and jid:
                        self._active_job_id = jid
                elif t == MSG_TRACE:
                    self._append_log_line(str(msg.get("line", "")), "trace")

                elif t == "done":
                    self._running = False
                    self._start_btn.configure(state="normal")
                    self._stop_btn.configure(state="disabled")
                    return

        except queue.Empty:
            pass
        if not self._closing and self.winfo_exists():
            self._queue_poll_after_id = self.after(80, self._poll_queue)

    # ------------------------------------------------------------------ #
    #  Save result
    # ------------------------------------------------------------------ #

    def _on_save_result(self) -> None:
        if self._result_type == "table" and self._last_result_df is not None:
            p = ctk.filedialog.asksaveasfilename(
                defaultextension=".csv",
                filetypes=[("CSV", "*.csv"), ("Excel", "*.xlsx"), ("Все", "*.*")],
            )
            if p:
                path = Path(p)
                if path.suffix.lower() == ".xlsx":
                    self._last_result_df.to_excel(path, index=False)
                else:
                    self._last_result_df.to_csv(path, index=False)
                self._set_status(f"Сохранено: {path}")
        else:
            p = ctk.filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[
                    ("Text", "*.txt"),
                    ("Markdown", "*.md"),
                    ("DOCX", "*.docx"),
                    ("PDF", "*.pdf"),
                    ("Все", "*.*"),
                ],
            )
            if p:
                out_path = Path(p)
                ext = out_path.suffix.lower()
                text = self._last_result_text or ""
                if ext == ".docx":
                    try:
                        self._export_to_docx(out_path, text)
                    except Exception as exc:
                        self._set_status(f"DOCX экспорт недоступен: {sanitize_for_log(str(exc))}", _STATUS_ERR)
                        return
                elif ext == ".pdf":
                    try:
                        self._export_to_pdf(out_path, text)
                    except Exception as exc:
                        self._set_status(f"PDF экспорт недоступен: {sanitize_for_log(str(exc))}", _STATUS_ERR)
                        return
                else:
                    out_path.write_text(text, encoding="utf-8")
                self._set_status(f"Сохранено: {out_path}")

    def _export_to_docx(self, path: Path, text: str) -> None:
        from docx import Document  # type: ignore[import-not-found]

        doc = Document()
        for line in text.splitlines():
            if line.startswith("## "):
                doc.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:].strip(), level=1)
            else:
                doc.add_paragraph(line)
        doc.save(path)

    def _export_to_pdf(self, path: Path, text: str) -> None:
        from weasyprint import HTML  # type: ignore[import-not-found]

        body = (
            text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace("\n", "<br>")
        )
        html = f"<html><meta charset='utf-8'><body style='font-family: DejaVu Sans, Arial, sans-serif;'>{body}</body></html>"
        HTML(string=html).write_pdf(str(path))

    def _update_preflight_label(self) -> None:
        model = self._model_var.get().strip()
        ctx = self._model_ctx.get(model, CONTEXT_FALLBACK)
        reserve = self._get_response_reserve(ctx)
        query = self._query_text.get("1.0", "end").strip() or "(запрос)"
        try:
            chunk = compute_dynamic_chunk_size(ctx, SYSTEM_PROMPT_MAP, query, response_reserve=reserve)
        except Exception:
            chunk = 0
        scout = "on" if self._scout_var.get() else "off"
        path = self._folder_path or self._file_path
        path_s = str(path) if path else "—"
        plan_line = ""
        if path and path.exists():
            try:
                from corpus_planner import format_plan_ru, plan_corpus

                plan = plan_corpus(
                    path,
                    query,
                    max(chunk, 500),
                    scout_mode=bool(self._scout_var.get()),
                    scout_threshold=float(self._scout_threshold_var.get() or "0.35"),
                )
                plan_line = format_plan_ru(plan)
            except Exception:
                plan_line = ""
        base = (
            f"Preflight: ctx={ctx} chunk≈{chunk} scout={scout} workers={self._workers_var.get()} | {path_s}"
        )
        text = f"{base} || {plan_line}" if plan_line else base
        self._preflight_label.configure(text=text.replace(",", " "), text_color=_md3.ON_SURFACE_VARIANT)

    def _maybe_first_run_wizard(self) -> None:
        from first_run import is_first_run, mark_first_run_complete

        if not is_first_run():
            return
        dlg = ctk.CTkToplevel(self)
        dlg.title("Первый запуск")
        dlg.geometry("420x220")
        ctk.CTkLabel(dlg, text="Настройте LM Studio", font=ctk.CTkFont(size=14, weight="bold")).pack(pady=12)
        ctk.CTkLabel(dlg, text="Укажите URL и API key, затем нажмите «Обновить модели».").pack(pady=6)

        def done() -> None:
            mark_first_run_complete({"base_url": self._url_var.get().strip()})
            dlg.destroy()
            self._on_fetch_models()

        ctk.CTkButton(dlg, text="Понятно", command=done).pack(pady=12)

    def _export_result(self, ext: str) -> None:
        if not self._last_result_text:
            self._set_status("Нет результата для экспорта", _STATUS_WARN)
            return
        p = ctk.filedialog.asksaveasfilename(defaultextension=ext, filetypes=[("All", "*.*")])
        if p:
            Path(p).write_text(self._last_result_text, encoding="utf-8")
            self._set_status(f"Экспорт: {p}")

    def _on_export_evidence_json(self) -> None:
        import json

        text = self._last_result_text or ""
        rows: list[dict[str, str]] = []
        for line in text.splitlines():
            if "|" in line and not line.strip().startswith("|--"):
                parts = [p.strip() for p in line.split("|") if p.strip()]
                if len(parts) >= 3:
                    rows.append({"file": parts[0], "chunk": parts[1], "quote": parts[2]})
        p = ctk.filedialog.asksaveasfilename(defaultextension=".json", filetypes=[("JSON", "*.json")])
        if p:
            Path(p).write_text(json.dumps({"evidence_matrix": rows}, ensure_ascii=False, indent=2), encoding="utf-8")
            self._set_status(f"Evidence JSON: {p}")

    def _on_show_run_history(self) -> None:
        from metrics import list_recent_runs

        runs = list_recent_runs(15)
        dlg = ctk.CTkToplevel(self)
        dlg.title("История запусков")
        dlg.geometry("640x320")
        tb = ctk.CTkTextbox(dlg, wrap="word")
        tb.pack(fill="both", expand=True, padx=10, pady=10)
        if not runs:
            tb.insert("1.0", "Пока нет записей метрик.")
        else:
            for r in runs:
                tb.insert(
                    "end",
                    f"#{r['id']} job={r.get('job_id','')} "
                    f"ok={r.get('chunks_ok')} fail={r.get('chunks_failed')} "
                    f"scout_skip={r.get('scout_skipped')} "
                    f"t={r.get('duration_s')}s\n  {r.get('query_preview','')}\n\n",
                )

    def _set_status(self, text: str, color: str = "gray") -> None:
        self._status_label.configure(text=text, text_color=color)


# ------------------------------------------------------------------ #
#  Worker functions (run in daemon threads)
# ------------------------------------------------------------------ #

def _run_processing(
    file_path: Path | None,
    folder_path: Path | None,
    query: str,
    rc: RunConfig,
    out_queue: queue.Queue[dict[str, Any]],
    stop_flag: Any = None,
) -> None:
    # Распаковка единой конфигурации прогона в локальные имена —
    # тело функции ниже работает с ними как прежде.
    model = rc.chat_model
    base_url = rc.base_url
    api_key = rc.api_key
    context_budget = rc.context_budget
    response_reserve = rc.response_reserve
    workers = rc.workers
    vision_model = rc.vision_model
    composer_model = rc.composer_model
    api_mode = rc.api_mode
    low_vram_mode = rc.low_vram_mode
    dual_instance_mode = False
    scout_mode = rc.scout_mode
    scout_relevance_threshold = rc.scout_threshold
    scout_model = rc.scout_model

    source_path = str(file_path or folder_path or "")
    dynamic_chunk_size = compute_dynamic_chunk_size(
        context_budget, SYSTEM_PROMPT_MAP, query, response_reserve=response_reserve,
    )
    out_queue.put({
        "type": MSG_TRACE,
        "line": (
            f"[PIPELINE] context_budget={context_budget} response_reserve={response_reserve} "
            f"dynamic_chunk_size={dynamic_chunk_size} workers={workers} "
            f"vision_model={vision_model or model} composer_model={composer_model or model} "
            f"api_mode={api_mode} low_vram={low_vram_mode} dual_instance={dual_instance_mode} "
            f"scout={scout_mode} scout_threshold={scout_relevance_threshold}"
        ),
    })
    q_preview = " ".join(query.strip().split())
    if len(q_preview) > 240:
        q_preview = q_preview[:240] + "..."
    out_queue.put({"type": MSG_TRACE, "line": f"[QUERY] {q_preview}"})

    def put_progress(
        current: int,
        total: int,
        phase: str = "map",
        from_cache: int = 0,
        **extra: Any,
    ) -> None:
        payload = {
            "type": MSG_PROGRESS, "current": current, "total": total,
            "phase": phase, "from_cache": from_cache,
        }
        payload.update(extra)
        out_queue.put(payload)

    # ---- archive as corpus (ZIP/TAR → per-file MAP, not one merged blob) ----
    if file_path is not None:
        from large_corpus_io import is_archive, corpus_input_root

        if is_archive(file_path):
            try:
                with corpus_input_root(file_path) as archive_root:
                    _run_folder_batch(
                        folder_path=archive_root,
                        query=query,
                        rc=rc,
                        dynamic_chunk_size=dynamic_chunk_size,
                        out_queue=out_queue,
                        put_progress=put_progress,
                        stop_flag=stop_flag,
                        job_id_root=file_path,
                        source_path=source_path,
                    )
            except ParseError as exc:
                out_queue.put({"type": MSG_ERROR, "message": sanitize_for_log(str(exc))})
            return

    # ---- folder → Map-Reduce over ALL files ----
    if folder_path is not None:
        _run_folder_batch(
            folder_path=folder_path,
            query=query,
            rc=rc,
            dynamic_chunk_size=dynamic_chunk_size,
            out_queue=out_queue,
            put_progress=put_progress,
            stop_flag=stop_flag,
            job_id_root=folder_path,
            source_path=source_path,
        )
        return

    # ---- single file ----
    assert file_path is not None
    try:
        kind, payload, _ = parse_file(file_path, dynamic_chunk_size, overlap_tokens=200, root_dir=file_path.parent)
    except ParseError as exc:
        out_queue.put({"type": MSG_ERROR, "message": sanitize_for_log(str(exc))})
        return
    except Exception as exc:
        logger.exception("parse_file failed")
        out_queue.put({"type": MSG_ERROR, "message": sanitize_for_log(str(exc))})
        return

    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        if kind == "text":
            chunks: list[str] = payload  # type: ignore[assignment]
            job_id = compute_job_id(
                file_path, query,
                chunk_size=dynamic_chunk_size,
                model=model,
                composer_model=composer_model,
            )
            out_queue.put({"type": MSG_JOB_ID, "job_id": job_id})
            result = loop.run_until_complete(
                run_map_reduce(
                    chunks=chunks,
                    user_query=query,
                    base_url=base_url,
                    api_key=api_key,
                    model=model,
                    workers=workers,
                    dynamic_chunk_size=dynamic_chunk_size,
                    on_progress=put_progress,
                    job_id=job_id,
                    max_context_tokens=context_budget,
                    composer_model=composer_model,
                    vision_model=vision_model,
                    api_mode=api_mode,
                    low_vram_mode=low_vram_mode,
                    dual_instance_mode=dual_instance_mode,
                    scout_mode=scout_mode,
                    scout_relevance_threshold=scout_relevance_threshold,
                    scout_model=scout_model,
                    source_path=source_path,
                    stop_flag=stop_flag,
                )
            )
            out_queue.put({"type": MSG_RESULT, "text": result})
        elif kind == "vision":
            vchunks: list[str] = payload  # type: ignore[assignment]
            job_id = compute_job_id(
                file_path, query,
                chunk_size=dynamic_chunk_size,
                model=model,
                composer_model=composer_model,
            )
            out_queue.put({"type": MSG_JOB_ID, "job_id": job_id})
            result = loop.run_until_complete(
                run_map_reduce(
                    chunks=vchunks,
                    user_query=query,
                    base_url=base_url,
                    api_key=api_key,
                    model=model,
                    workers=workers,
                    dynamic_chunk_size=dynamic_chunk_size,
                    on_progress=put_progress,
                    job_id=job_id,
                    max_context_tokens=context_budget,
                    composer_model=composer_model,
                    vision_model=vision_model,
                    api_mode=api_mode,
                    low_vram_mode=low_vram_mode,
                    dual_instance_mode=dual_instance_mode,
                    scout_mode=scout_mode,
                    scout_relevance_threshold=scout_relevance_threshold,
                    scout_model=scout_model,
                    source_path=source_path,
                    stop_flag=stop_flag,
                )
            )
            out_queue.put({"type": MSG_RESULT, "text": result})
        else:
            batches = payload
            result_df = loop.run_until_complete(
                run_batching(
                    batches=batches,
                    user_query=query,
                    base_url=base_url,
                    api_key=api_key,
                    model=model,
                    workers=workers,
                    on_progress=lambda c, t: put_progress(c, t, "batch"),
                    api_mode=api_mode,
                )
            )
            saved = None
            if not result_df.empty:
                try:
                    out_path = file_path.parent / f"{file_path.stem}_result.csv"
                    result_df.to_csv(out_path, index=False, encoding="utf-8")
                    saved = str(out_path)
                except Exception as exc:
                    logger.warning("Cannot save result table: %s", exc)
            out_queue.put({"type": MSG_RESULT_DF, "df": result_df, "saved_path": saved})
    finally:
        loop.close()


def _run_folder_batch(
    folder_path: Path,
    query: str,
    rc: RunConfig,
    dynamic_chunk_size: int,
    out_queue: queue.Queue[dict[str, Any]],
    put_progress: Any,
    stop_flag: Any,
    job_id_root: Path | None = None,
    source_path: str = "",
) -> None:
    from pipeline import _iter_files, _to_chunks
    from corpus_planner import filter_files_by_relevance
    from chunk_store import ChunkStore

    # Распаковка конфигурации прогона в локальные имена.
    model = rc.chat_model
    base_url = rc.base_url
    api_key = rc.api_key
    context_budget = rc.context_budget
    workers = rc.workers
    vision_model = rc.vision_model
    composer_model = rc.composer_model
    api_mode = rc.api_mode
    low_vram_mode = rc.low_vram_mode
    dual_instance_mode = False
    scout_mode = rc.scout_mode
    scout_relevance_threshold = rc.scout_threshold
    scout_model = rc.scout_model

    all_files = _iter_files([folder_path])
    if not all_files:
        out_queue.put({"type": MSG_ERROR,
                       "message": "Нет поддерживаемых файлов в папке"})
        return

    filtered_files, files_skipped = filter_files_by_relevance(
        all_files,
        query,
        scout_mode=scout_mode,
        threshold=scout_relevance_threshold,
    )
    if files_skipped:
        out_queue.put({
            "type": MSG_TRACE,
            "line": f"[FILE SCOUT] skipped {files_skipped} low-relevance files (heuristic)",
        })

    id_root = job_id_root or folder_path
    job_id = compute_job_id(
        id_root, query,
        file_paths=filtered_files,
        chunk_size=dynamic_chunk_size,
        model=model,
        composer_model=composer_model,
    )
    corpus_src = source_path or str(id_root)
    out_queue.put({"type": MSG_JOB_ID, "job_id": job_id})
    chunk_store = ChunkStore(job_id)
    try:
        extract_workers = min(8, os.cpu_count() or 4)
        completed_files = 0
        chunks_by_file: dict[int, list[str]] = {}

        def _extract_one(idx_fp: tuple[int, Path]) -> tuple[int, list[str], str]:
            idx, fp = idx_fp
            file_chunks = [
                dc.text for dc in _to_chunks(fp, dynamic_chunk_size, 200, root_dir=folder_path)
            ]
            try:
                rel = str(fp.relative_to(folder_path)).replace("\\", "/")
            except ValueError:
                rel = fp.name
            return idx, file_chunks, rel

        with ThreadPoolExecutor(max_workers=extract_workers) as pool:
            futures = {
                pool.submit(_extract_one, (i, fp)): (i, fp)
                for i, fp in enumerate(filtered_files)
            }
            for fut in as_completed(futures):
                if stop_flag and stop_flag():
                    out_queue.put({"type": MSG_ERROR, "message": "Остановлено пользователем"})
                    return
                try:
                    idx, file_chunks, display_rel = fut.result()
                except Exception as exc:
                    logger.warning("Extract failed for file: %s", exc)
                    idx, file_chunks, display_rel = futures[fut][0], [], "unknown"
                chunks_by_file[idx] = file_chunks
                completed_files += 1
                put_progress(completed_files, len(filtered_files), "extract", file=display_rel)

        for i in range(len(filtered_files)):
            chunk_store.extend(chunks_by_file.get(i, []))

        if len(chunk_store) == 0:
            out_queue.put({"type": MSG_ERROR,
                           "message": "Не удалось извлечь текст из файлов"})
            return

        all_chunks: list[str] | ChunkStore = chunk_store
        logger.info(
            "Folder batch: files=%s filtered=%s chunks=%s spilled=%s",
            len(all_files),
            len(filtered_files),
            len(chunk_store),
            chunk_store._spilled,
        )

        try:
            from corpus_manifest import build_corpus_manifest, manifest_to_json

            manifest = build_corpus_manifest([folder_path])
            # Манифест пишем в кэш-директорию, а НЕ в исходную папку:
            # иначе файл попадёт в корпус на следующем прогоне и сдвинет
            # corpus_fingerprint → нестабильный job_id → resume не работает.
            from cache import CACHE_DIR

            manifest_dir = CACHE_DIR / "manifests"
            manifest_dir.mkdir(parents=True, exist_ok=True)
            manifest_path = manifest_dir / f"{job_id}.json"
            manifest_path.write_text(manifest_to_json(manifest), encoding="utf-8")
            out_queue.put({
                "type": MSG_TRACE,
                "line": f"[MANIFEST] files={manifest.get('files_total')} bytes={manifest.get('total_bytes')}",
            })
        except Exception as exc:
            logger.warning("Manifest build failed: %s", exc)

        _relevant_files_count: list[int] = [0]
        _orig_put_progress = put_progress

        def _tracking_put_progress(
            current: int,
            total: int,
            phase: str = "map",
            from_cache: int = 0,
            **extra: Any,
        ) -> None:
            if phase == "map_relevant":
                _relevant_files_count[0] = extra.get("relevant_files", current)
            _orig_put_progress(current, total, phase, from_cache, **extra)

        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            result = loop.run_until_complete(
                run_map_reduce(
                    chunks=all_chunks,  # type: ignore[arg-type]
                    user_query=query,
                    base_url=base_url,
                    api_key=api_key,
                    model=model,
                    workers=workers,
                    dynamic_chunk_size=dynamic_chunk_size,
                    on_progress=_tracking_put_progress,
                    job_id=job_id,
                    max_context_tokens=context_budget,
                    composer_model=composer_model,
                    vision_model=vision_model,
                    api_mode=api_mode,
                    low_vram_mode=low_vram_mode,
                    dual_instance_mode=dual_instance_mode,
                    scout_mode=scout_mode,
                    scout_relevance_threshold=scout_relevance_threshold,
                    scout_model=scout_model,
                    source_path=corpus_src,
                    stop_flag=stop_flag,
                )
            )
        finally:
            loop.close()

        relevant_count = _relevant_files_count[0]
        relevant_suffix = f"  |  релевантных файлов: {relevant_count}" if relevant_count > 0 else ""
        out_queue.put({
            "type": MSG_RESULT,
            "text": (
                f"Обработано: {len(filtered_files)} файлов (из {len(all_files)})  |  "
                f"{len(all_chunks)} чанков{relevant_suffix}\n"
                f"{'─' * 60}\n\n{result}"
            ),
        })
    finally:
        chunk_store.cleanup()
