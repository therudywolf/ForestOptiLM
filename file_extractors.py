# SPDX-License-Identifier: AGPL-3.0-or-later
# Copyright (C) 2025 therudywolf <https://github.com/therudywolf>
#
# This file is part of ForestOptiLM / Nocturne Data Forge.
# ForestOptiLM is free software: you can redistribute it and/or modify it
# under the terms of the GNU Affero General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# ForestOptiLM is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY
# or FITNESS FOR A PARTICULAR PURPOSE. See the GNU Affero General Public
# License for more details.
#
# You should have received a copy of the GNU Affero General Public License
# along with ForestOptiLM. If not, see <https://www.gnu.org/licenses/>.
"""
Nocturne Data Forge — всеядный маршрутизатор форматов.
Поддерживает: текст, код, таблицы, архивы, документы.
"""
from __future__ import annotations

import gzip
import io
import json
import logging
import os
import re
import tarfile
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Callable, Literal

import pandas as pd
from pandas.errors import EmptyDataError

logger = logging.getLogger("nocturne")

ContentKind = Literal["text", "table", "vision"]

# Изображения для vision-модели (v1)
IMAGE_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff",
}

# ------------------------------------------------------------------ #
#  Optional heavy dependencies
# ------------------------------------------------------------------ #

try:
    import pdfplumber
except ImportError:
    pdfplumber = None  # type: ignore

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None  # type: ignore

try:
    import yaml
except ImportError:
    yaml = None  # type: ignore

try:
    from odf import text as odf_text
    from odf.opendocument import load as odf_load
except ImportError:
    odf_text = None  # type: ignore
    odf_load = None  # type: ignore

try:
    import ebooklib
    from ebooklib import epub
except ImportError:
    epub = None  # type: ignore

try:
    from striprtf.striprtf import rtf_to_text
except ImportError:
    rtf_to_text = None  # type: ignore


class ParseError(Exception):
    pass


# ------------------------------------------------------------------ #
#  Encoding helpers
# ------------------------------------------------------------------ #

_ENCODINGS = ("utf-8-sig", "utf-8", "cp1251", "cp1252", "latin-1")


def _decode(raw: bytes) -> str:
    # Detect UTF-16 BOM
    if raw[:2] in (b"\xff\xfe", b"\xfe\xff"):
        return raw.decode("utf-16", errors="replace")
    for enc in _ENCODINGS:
        try:
            return raw.decode(enc)
        except (UnicodeDecodeError, LookupError):
            continue
    return raw.decode("utf-8", errors="replace")


def _read_plain_text(path: Path) -> str:
    return _decode(path.read_bytes())


# ------------------------------------------------------------------ #
#  Text extractors
# ------------------------------------------------------------------ #

def _read_rtf(path: Path) -> str:
    if rtf_to_text is None:
        raise ParseError("striprtf is not installed")
    return rtf_to_text(_decode(path.read_bytes()))


def _read_pdf(path: Path) -> str:
    if pdfplumber is None:
        raise ParseError("pdfplumber is not installed")
    parts: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)


def _read_docx(path: Path) -> str:
    if DocxDocument is None:
        raise ParseError("python-docx is not installed")
    doc = DocxDocument(path)
    lines: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text)
    # Also extract tables from docx
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(lines)


def _read_odt(path: Path) -> str:
    if odf_load is None or odf_text is None:
        raise ParseError("odfpy is not installed")
    doc = odf_load(str(path))
    parts: list[str] = []
    for el in doc.getElementsByType(odf_text.P):
        s = getattr(el, "getTextContent", None)
        parts.append(s() if callable(s) else str(el))
    return "\n".join(parts)


def _read_epub(path: Path) -> str:
    if epub is None:
        raise ParseError("ebooklib is not installed")
    book = epub.read_epub(str(path))
    parts: list[str] = []
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            try:
                html = item.get_content().decode("utf-8", errors="replace")
                if BeautifulSoup:
                    parts.append(BeautifulSoup(html, "html.parser").get_text(separator="\n"))
                else:
                    parts.append(re.sub(r"<[^>]+>", " ", html))
            except Exception as exc:
                logger.debug("epub item %s skip: %s", item.get_name(), exc)
    return "\n".join(parts)


def _read_html_text(path: Path) -> str:
    raw = path.read_bytes()
    text = _decode(raw)
    if BeautifulSoup:
        return BeautifulSoup(text, "html.parser").get_text(separator="\n", strip=True)
    return re.sub(r"<[^>]+>", " ", text)


# ------------------------------------------------------------------ #
#  Table extractors
# ------------------------------------------------------------------ #

def _read_csv(path: Path) -> pd.DataFrame:
    raw = path.read_bytes()
    # Detect separator heuristically
    sample = _decode(raw[:4096])
    sep_scores = {
        "\t": sample.count("\t"),
        ";": sample.count(";"),
        ",": sample.count(","),
    }
    sep = max(sep_scores.items(), key=lambda kv: kv[1])[0]
    for enc in _ENCODINGS:
        try:
            df = pd.read_csv(
                io.BytesIO(raw), encoding=enc, sep=sep,
                on_bad_lines="skip", low_memory=False,
            )
            if not df.empty:
                return df
        except (UnicodeDecodeError, EmptyDataError):
            continue
        except Exception:
            continue
    return pd.DataFrame()


def _read_xlsx(path: Path) -> pd.DataFrame:
    return pd.read_excel(str(path), engine="openpyxl")


def _read_xls(path: Path) -> pd.DataFrame:
    return pd.read_excel(str(path), engine="xlrd")


def _read_json_table(path: Path) -> pd.DataFrame:
    try:
        data = json.loads(_decode(path.read_bytes()))
        if isinstance(data, list):
            if data and isinstance(data[0], dict):
                return pd.json_normalize(data)
            return pd.DataFrame({"value": data})
        if isinstance(data, dict):
            return pd.json_normalize([data])
        return pd.DataFrame()
    except Exception as exc:
        logger.debug("JSON table parse failed %s: %s", path, exc)
        return pd.DataFrame()


def _read_yaml_table(path: Path) -> pd.DataFrame:
    if yaml is None:
        raise ParseError("pyyaml is not installed")
    data = yaml.safe_load(_read_plain_text(path))
    if data is None:
        return pd.DataFrame()
    if isinstance(data, list) and data and isinstance(data[0], dict):
        return pd.DataFrame(data)
    if isinstance(data, dict):
        return pd.DataFrame([data])
    return pd.DataFrame()


def _df_to_text(path: Path, df: pd.DataFrame) -> str:
    """Convert large or complex DataFrames to readable text for LLM."""
    cols = ", ".join(str(c) for c in df.columns)
    total = len(df)
    # For very large tables only show first 200 rows as text
    preview_rows = df.head(200)
    return (
        f"Таблица: {path.name}  |  строк: {total}  |  колонки: {cols}\n"
        + preview_rows.to_string(index=False, max_colwidth=120)
    )


# ------------------------------------------------------------------ #
#  Registries
# ------------------------------------------------------------------ #

TEXT_EXTRACTORS: dict[str, Callable[[Path], str]] = {
    # Plain text / markup
    ".txt":  _read_plain_text,
    ".md":   _read_plain_text,
    ".rst":  _read_plain_text,
    ".log":  _read_plain_text,
    # Config
    ".ini":  _read_plain_text,
    ".cfg":  _read_plain_text,
    ".conf": _read_plain_text,
    ".toml": _read_plain_text,
    ".env":  _read_plain_text,
    ".properties": _read_plain_text,
    # Rich documents
    ".rtf":  _read_rtf,
    ".pdf":  _read_pdf,
    ".docx": _read_docx,
    ".odt":  _read_odt,
    ".epub": _read_epub,
    # Web
    ".html": _read_html_text,
    ".htm":  _read_html_text,
    # Code – all treated as plain text
    ".py":   _read_plain_text,
    ".js":   _read_plain_text,
    ".ts":   _read_plain_text,
    ".tsx":  _read_plain_text,
    ".jsx":  _read_plain_text,
    ".c":    _read_plain_text,
    ".cpp":  _read_plain_text,
    ".cc":   _read_plain_text,
    ".h":    _read_plain_text,
    ".hpp":  _read_plain_text,
    ".java": _read_plain_text,
    ".kt":   _read_plain_text,
    ".scala": _read_plain_text,
    ".go":   _read_plain_text,
    ".rs":   _read_plain_text,
    ".rb":   _read_plain_text,
    ".php":  _read_plain_text,
    ".cs":   _read_plain_text,
    ".swift": _read_plain_text,
    ".sql":  _read_plain_text,
    ".sh":   _read_plain_text,
    ".bash": _read_plain_text,
    ".bat":  _read_plain_text,
    ".cmd":  _read_plain_text,
    ".ps1":  _read_plain_text,
    ".r":    _read_plain_text,
    ".lua":  _read_plain_text,
    ".dart": _read_plain_text,
    # Markup data — kept consistent with large_corpus_io.STREAMING_PLAIN_SUFFIXES
    ".xml":  _read_plain_text,
}

TABLE_EXTRACTORS: dict[str, Callable[[Path], pd.DataFrame]] = {
    ".csv":  _read_csv,
    ".xlsx": _read_xlsx,
    ".xls":  _read_xls,
    ".json": _read_json_table,
    ".yaml": _read_yaml_table,
    ".yml":  _read_yaml_table,
}

ARCHIVE_EXTENSIONS = {".zip", ".tar", ".gz", ".tgz"}

# Extensions to skip silently (binary artifacts, noise)
_SKIP_EXTENSIONS = {
    ".class", ".jar", ".war", ".ear",
    ".pyc", ".pyo", ".pyd",
    ".exe", ".dll", ".so", ".dylib",
    ".ico", ".svg",
    ".mp3", ".mp4", ".avi", ".mov", ".woff", ".woff2", ".ttf", ".eot",
    ".bin", ".dat",
}


# ------------------------------------------------------------------ #
#  Core extraction logic
# ------------------------------------------------------------------ #

def _is_archive(path: Path) -> bool:
    name = path.name.lower()
    if name.endswith(".tar.gz") or name.endswith(".tgz"):
        return True
    return path.suffix.lower() in {".zip", ".tar", ".gz"}


def _is_probably_binary(path: Path) -> bool:
    try:
        raw = path.read_bytes()[:8192]
    except Exception:
        return True
    if not raw:
        return False
    # UTF-16 BOM → not binary (text)
    if raw[:2] in (b"\xff\xfe", b"\xfe\xff"):
        return False
    if b"\x00" in raw:
        return True
    printable = sum((32 <= b <= 126) or b in (9, 10, 13) for b in raw)
    return (printable / len(raw)) < 0.65


def _extract_single_file(path: Path) -> tuple[ContentKind, str | pd.DataFrame]:
    suffix = path.suffix.lower()
    if path.name.lower().endswith(".tar.gz"):
        suffix = ".tar.gz"

    if suffix in _SKIP_EXTENSIONS:
        logger.debug("Skip by extension: %s", path)
        raise ParseError(f"Extension {suffix} skipped by design")

    if suffix == ".doc":
        raise ParseError(".doc not supported; convert to .docx")

    if suffix in IMAGE_EXTENSIONS:
        return ("vision", str(path.resolve()))

    # Try table extractor first
    if suffix in TABLE_EXTRACTORS:
        try:
            df = TABLE_EXTRACTORS[suffix](path)
            if df.empty:
                # Fall through to text
                pass
            elif len(df) > 2000:
                # Very large table → represent as text summary for LLM
                return ("text", _df_to_text(path, df))
            else:
                return ("table", df)
        except Exception as exc:
            logger.debug("Table extract failed %s: %s; try text", path, exc)

    if suffix in TEXT_EXTRACTORS:
        try:
            return ("text", TEXT_EXTRACTORS[suffix](path))
        except ParseError:
            raise
        except Exception as exc:
            raise ParseError(f"Text extract failed: {exc}") from exc

    if _is_probably_binary(path):
        logger.debug("Skip binary: %s", path)
        raise ParseError("Binary file skipped")

    # Unknown extension but looks like text → read as plain
    try:
        return ("text", _read_plain_text(path))
    except Exception as exc:
        raise ParseError(f"Cannot read unknown file: {exc}") from exc


# Streaming copy chunk size for .gz extraction.
_GZ_CHUNK = 1024 * 1024  # 1 MiB


def _max_uncompressed_bytes() -> int:
    """Cap on total uncompressed archive payload (decompression-bomb guard).

    Read directly from the environment to avoid a circular import with
    large_corpus_io (which imports this module). A value of 0 disables the limit.
    """
    raw = os.getenv("NOCTURNE_MAX_UNCOMPRESSED_BYTES", "8589934592").strip()  # 8 GiB
    try:
        return max(0, int(raw))
    except ValueError:
        return 8 * 1024 * 1024 * 1024


def _is_within(base: Path, target: Path) -> bool:
    """True iff ``target`` is ``base`` itself or nested inside it.

    Uses os.path.commonpath on resolved, normalized paths so that a sibling
    such as ``/tmp/foobar`` is NOT considered inside ``/tmp/foo``. Works on
    Windows (case-insensitive, drive-aware): commonpath raises ValueError for
    paths on different drives, which is treated as "not within".
    """
    try:
        base_r = base.resolve()
        target_r = target.resolve()
    except (OSError, RuntimeError):
        return False
    try:
        return os.path.commonpath([str(base_r), str(target_r)]) == str(base_r)
    except ValueError:
        # Different drives / mix of absolute and relative — cannot be within.
        return False


def _extract_archive_to_dir(archive_path: Path, destination: Path) -> None:
    name = archive_path.name.lower()
    suf = archive_path.suffix.lower()
    dest_root = destination.resolve()
    limit = _max_uncompressed_bytes()
    try:
        if suf == ".zip":
            with zipfile.ZipFile(archive_path, "r") as zf:
                total = 0
                for member in zf.infolist():
                    target = destination / member.filename
                    if not _is_within(dest_root, target):
                        logger.warning(
                            "Skip suspicious zip member outside target dir: %s",
                            member.filename,
                        )
                        continue
                    total += int(member.file_size)
                    if limit and total > limit:
                        raise ParseError(
                            f"Archive uncompressed size exceeds "
                            f"NOCTURNE_MAX_UNCOMPRESSED_BYTES ({limit}): {archive_path}",
                        )
                    zf.extract(member, destination)
        elif name.endswith(".tar.gz") or suf in {".tar", ".tgz"}:
            with tarfile.open(archive_path, "r:*") as tf:
                total = 0
                for member in tf.getmembers():
                    if member.issym() or member.islnk():
                        logger.warning(
                            "Skip link member in tar (potential escape): %s",
                            member.name,
                        )
                        continue
                    if member.isdev():
                        logger.warning("Skip device member in tar: %s", member.name)
                        continue
                    member_name = member.name
                    if os.path.isabs(member_name) or member_name.startswith(("/", "\\")):
                        logger.warning(
                            "Skip tar member with absolute name: %s", member_name,
                        )
                        continue
                    target = destination / member_name
                    if not _is_within(dest_root, target):
                        logger.warning(
                            "Skip suspicious tar member outside target dir: %s",
                            member_name,
                        )
                        continue
                    total += int(member.size)
                    if limit and total > limit:
                        raise ParseError(
                            f"Archive uncompressed size exceeds "
                            f"NOCTURNE_MAX_UNCOMPRESSED_BYTES ({limit}): {archive_path}",
                        )
                    tf.extract(member, destination)
        elif suf == ".gz":
            out_file = destination / archive_path.stem
            written = 0
            with gzip.open(archive_path, "rb") as gz, out_file.open("wb") as out:
                while True:
                    chunk = gz.read(_GZ_CHUNK)
                    if not chunk:
                        break
                    written += len(chunk)
                    if limit and written > limit:
                        raise ParseError(
                            f"Gzip uncompressed size exceeds "
                            f"NOCTURNE_MAX_UNCOMPRESSED_BYTES ({limit}): {archive_path}",
                        )
                    out.write(chunk)
        else:
            raise ParseError(f"Unsupported archive type: {archive_path}")
    except ParseError:
        raise
    except Exception as exc:
        logger.warning("Archive extract failed %s: %s", archive_path, exc)
        raise ParseError(f"Cannot extract: {exc}") from exc


def _collect_from_path(path: Path, root: Path, parts: list[str]) -> None:
    if path.is_dir():
        for child in sorted(path.iterdir()):
            _collect_from_path(child, root, parts)
        return

    if _is_archive(path):
        try:
            with TemporaryDirectory(prefix="nocturne_nested_") as tmp:
                tmp_p = Path(tmp)
                _extract_archive_to_dir(path, tmp_p)
                _collect_from_path(tmp_p, root, parts)
        except Exception as exc:
            logger.debug("Skip nested archive %s: %s", path, exc)
        return

    try:
        kind, content = _extract_single_file(path)
        try:
            rel = path.relative_to(root)
        except ValueError:
            rel = path.name
        if kind == "text":
            text = str(content).strip()
            if text:
                parts.append(f"--- {rel} ---\n{text}")
        else:
            df: pd.DataFrame = content  # type: ignore[assignment]
            if not df.empty:
                parts.append(f"--- {rel} ---\n{_df_to_text(path, df)}")
    except Exception as exc:
        logger.debug("Skip %s: %s", path, exc)


def _expand_archive_to_text(path: Path) -> str:
    parts: list[str] = []
    with TemporaryDirectory(prefix="nocturne_") as tmp:
        tmp_p = Path(tmp)
        _extract_archive_to_dir(path, tmp_p)
        _collect_from_path(tmp_p, tmp_p, parts)
    return "\n\n".join(parts)


def extract_content(path: Path) -> tuple[ContentKind, str | pd.DataFrame]:
    """
    Facade: detect type → extract.
    Returns ("text", str) | ("table", pd.DataFrame) | ("vision", path_str).
    """
    path = Path(path)
    if not path.exists():
        raise ParseError(f"File not found: {path}")
    if _is_archive(path):
        text = _expand_archive_to_text(path)
        if not text.strip():
            raise ParseError(f"Archive had no readable content: {path}")
        return ("text", text)
    return _extract_single_file(path)


# ------------------------------------------------------------------ #
#  File Metadata (fast sampling — no full re-read)
# ------------------------------------------------------------------ #

_FORMAT_LABELS: dict[str, str] = {
    ".docx": "DOCX document",
    ".doc":  "DOC document",
    ".pdf":  "PDF document",
    ".odt":  "ODT document",
    ".epub": "EPUB document",
    ".rtf":  "RTF document",
    ".md":   "Markdown",
    ".txt":  "plain text",
    ".log":  "log file",
    ".rst":  "reStructuredText",
    ".html": "HTML document",
    ".htm":  "HTML document",
    ".csv":  "CSV table",
    ".xlsx": "Excel table",
    ".xls":  "Excel table",
    ".json": "JSON data",
    ".yaml": "YAML data",
    ".yml":  "YAML data",
    ".toml": "TOML config",
    ".ini":  "INI config",
    ".cfg":  "config file",
    ".properties": "properties file",
    ".py":   "Python source",
    ".js":   "JavaScript source",
    ".ts":   "TypeScript source",
    ".tsx":  "TypeScript/React source",
    ".java": "Java source",
    ".kt":   "Kotlin source",
    ".go":   "Go source",
    ".rs":   "Rust source",
    ".sql":  "SQL file",
    ".sh":   "shell script",
    ".bat":  "batch script",
    ".ps1":  "PowerShell script",
    ".xml":  "XML data",
    ".zip":  "ZIP archive",
    ".tar":  "TAR archive",
}


def _fast_title_from_text(path: Path) -> str:
    """Read first 4 KB of a text file and return the first non-empty heading or line."""
    try:
        raw = path.read_bytes()[:4096]
        text = _decode(raw)
        suffix = path.suffix.lower()
        if suffix in (".md", ".rst", ".txt", ".log"):
            for line in text.splitlines():
                stripped = line.strip()
                if stripped.startswith("#"):
                    return re.sub(r"^#+\s*", "", stripped).strip()
                if stripped:
                    return stripped[:120]
        elif suffix in (".html", ".htm"):
            m = re.search(r"<title[^>]*>([^<]+)</title>", text, re.IGNORECASE)
            if m:
                return m.group(1).strip()[:120]
            m = re.search(r"<h1[^>]*>([^<]+)</h1>", text, re.IGNORECASE)
            if m:
                return re.sub(r"<[^>]+>", "", m.group(1)).strip()[:120]
        elif suffix in (".py", ".js", ".ts", ".tsx", ".java", ".kt", ".go", ".rs", ".sql"):
            for line in text.splitlines():
                stripped = line.strip()
                # Docstring or module comment
                if stripped.startswith(("\"\"\"", "'''", "//", "#", "/*", "*")):
                    cleaned = re.sub(r'^["\'\s/*#]+', "", stripped).strip()
                    if len(cleaned) > 10:
                        return cleaned[:120]
    except Exception:
        pass
    return ""


def _fast_title_from_docx(path: Path) -> str:
    """Extract title from DOCX core properties or first heading paragraph."""
    if DocxDocument is None:
        return ""
    try:
        doc = DocxDocument(path)
        props = doc.core_properties
        if props.title and props.title.strip():
            return props.title.strip()[:120]
        # Fall back to first heading paragraph
        for p in doc.paragraphs:
            if p.style.name.startswith("Heading") and p.text.strip():
                return p.text.strip()[:120]
            if p.text.strip():
                return p.text.strip()[:120]
    except Exception:
        pass
    return ""


def _fast_title_from_pdf(path: Path) -> str:
    """Extract title from PDF metadata."""
    if pdfplumber is None:
        return ""
    try:
        with pdfplumber.open(path) as pdf:
            meta = pdf.metadata or {}
            title = meta.get("Title") or meta.get("title") or ""
            if title and title.strip():
                return str(title).strip()[:120]
            # Try first page first text block
            if pdf.pages:
                text = pdf.pages[0].extract_text() or ""
                for line in text.splitlines():
                    if line.strip():
                        return line.strip()[:120]
    except Exception:
        pass
    return ""


def _fast_title_from_csv(path: Path) -> str:
    """Return first non-empty header column from CSV — suggests what the file contains."""
    try:
        raw = path.read_bytes()[:2048]
        text = _decode(raw)
        lines = text.splitlines()
        if lines:
            # Try to detect separator
            first = lines[0]
            sep = ";" if first.count(";") > first.count(",") else ","
            cols = [c.strip().strip('"') for c in first.split(sep) if c.strip()]
            if cols:
                return ", ".join(cols[:5])
    except Exception:
        pass
    return ""


def _fast_title_from_xlsx(path: Path) -> str:
    """Return column names from first sheet of an Excel file."""
    try:
        df = pd.read_excel(path, nrows=0)
        cols = [str(c).strip() for c in df.columns if str(c).strip()]
        if cols:
            return ", ".join(cols[:5])
    except Exception:
        pass
    return ""


def _labels_from_path(path: Path, root_dir: Path | None = None) -> str:
    """
    Derive taxonomy labels from folder components of the path.
    Uses parts between root_dir and the filename, split on common separators.
    Example: reports/sonar/wildfly-certificates/scan.csv
      → "sonar, wildfly-certificates"
    """
    try:
        base = root_dir if root_dir is not None else path.parent
        try:
            rel = path.relative_to(base)
        except ValueError:
            rel = path
        parts = list(rel.parts[:-1])  # exclude filename
        if not parts:
            # use parent folder name
            parts = [path.parent.name]
        # Clean up: remove very generic single-char parts, dedup
        seen: set[str] = set()
        clean: list[str] = []
        for p in parts:
            p = p.strip()
            if p and p not in seen and len(p) > 1:
                seen.add(p)
                clean.append(p)
        return ", ".join(clean[:6])
    except Exception:
        return ""


def extract_file_metadata(path: Path, root_dir: Path | None = None) -> dict[str, str]:
    """
    Fast sampling of file metadata for chunk header enrichment.
    Returns a dict with keys: title, labels, format.
    Never raises — on any error returns partial/empty metadata.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    # Format label
    fmt = _FORMAT_LABELS.get(suffix, f"{suffix.lstrip('.').upper()} file" if suffix else "file")

    # Title — dispatcher
    title = ""
    try:
        if suffix == ".docx":
            title = _fast_title_from_docx(path)
        elif suffix == ".pdf":
            title = _fast_title_from_pdf(path)
        elif suffix in (".csv",):
            title = _fast_title_from_csv(path)
            fmt = "CSV table"
        elif suffix in (".xlsx", ".xls"):
            title = _fast_title_from_xlsx(path)
        elif suffix in (".md", ".txt", ".log", ".rst", ".html", ".htm",
                        ".py", ".js", ".ts", ".tsx", ".java", ".kt",
                        ".go", ".rs", ".sql", ".sh", ".bat", ".ps1"):
            title = _fast_title_from_text(path)
        elif suffix in (".json",):
            # For JSON, use filename stem as title — no quick way to get a title
            title = path.stem.replace("_", " ").replace("-", " ")
    except Exception:
        title = ""

    if not title:
        title = path.stem.replace("_", " ").replace("-", " ")

    labels = _labels_from_path(path, root_dir)

    return {
        "title": title[:200],
        "labels": labels,
        "format": fmt,
    }
