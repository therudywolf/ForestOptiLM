# SPDX-License-Identifier: AGPL-3.0-or-later
"""DOCX/XLSX извлечение: таблицы→Markdown, все листы, встроенные картинки→vision."""
from __future__ import annotations

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

import pandas as pd

import chunking
import file_extractors as fe

# Минимальный валидный 1×1 PNG (для проверки извлечения картинок из docx).
_PNG_1x1 = bytes.fromhex(
    "89504e470d0a1a0a0000000d4948445200000001000000010802000000"
    "907753de0000000c4944415408d763f8cfc0f01f0005000001ff5e9b9e"
    "0000000049454e44ae426082"
)


class TestRowsToMarkdown(unittest.TestCase):
    def test_header_separator_and_rows(self) -> None:
        md = fe._rows_to_markdown([["A", "B"], ["1", "2"], ["3", "4"]])
        lines = md.splitlines()
        self.assertEqual(lines[0], "| A | B |")
        self.assertEqual(lines[1], "| --- | --- |")
        self.assertEqual(lines[2], "| 1 | 2 |")

    def test_escapes_pipe_and_newline(self) -> None:
        md = fe._rows_to_markdown([["a|b", "c\nd"]])
        self.assertIn(r"a\|b", md)
        self.assertNotIn("\n", md.splitlines()[0])  # перенос внутри ячейки убран

    def test_ragged_rows_padded(self) -> None:
        md = fe._rows_to_markdown([["a", "b", "c"], ["x"]])
        self.assertEqual(md.splitlines()[-1], "| x |  |  |")

    def test_empty(self) -> None:
        self.assertEqual(fe._rows_to_markdown([]), "")
        self.assertEqual(fe._rows_to_markdown([["", "  "]]), "")


class TestRasterExt(unittest.TestCase):
    def test_known_formats(self) -> None:
        self.assertEqual(chunking._raster_ext(_PNG_1x1), ".png")
        self.assertEqual(chunking._raster_ext(b"\xff\xd8\xff\xe0junk"), ".jpg")
        self.assertEqual(chunking._raster_ext(b"GIF89a..."), ".gif")
        self.assertEqual(chunking._raster_ext(b"BMxx"), ".bmp")
        self.assertEqual(chunking._raster_ext(b"RIFF\x00\x00\x00\x00WEBP"), ".webp")

    def test_vector_skipped(self) -> None:
        # EMF/WMF (вектор) → пусто, vision их не обработает.
        self.assertEqual(chunking._raster_ext(b"\x01\x00\x00\x00 emf"), "")
        self.assertEqual(chunking._raster_ext(b""), "")


try:
    from PIL import Image as _PILImage  # noqa: F401
    _HAS_PIL = True
except Exception:
    _HAS_PIL = False


@unittest.skipIf(fe.DocxDocument is None or not _HAS_PIL, "python-docx/Pillow not installed")
class TestDocx(unittest.TestCase):
    def _make_docx(self, with_image: bool = False) -> Path:
        from docx import Document
        doc = Document()
        doc.add_paragraph("Обычный абзац текста.")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "ВМ"
        table.rows[0].cells[1].text = "Подсистема"
        table.rows[1].cells[0].text = "host-01"
        table.rows[1].cells[1].text = "Beta"
        if with_image:
            from io import BytesIO

            from PIL import Image
            buf = BytesIO()
            Image.new("RGB", (8, 8), (200, 40, 40)).save(buf, format="PNG")
            buf.seek(0)
            doc.add_picture(buf)
        p = Path(self._tmp.name) / ("img.docx" if with_image else "t.docx")
        doc.save(str(p))
        return p

    def setUp(self) -> None:
        self._tmp = TemporaryDirectory()
        self.addCleanup(self._tmp.cleanup)

    def test_table_rendered_as_markdown(self) -> None:
        text = fe._read_docx(self._make_docx())
        self.assertIn("Обычный абзац", text)
        self.assertIn("| ВМ | Подсистема |", text)
        self.assertIn("| --- | --- |", text)
        self.assertIn("| host-01 | Beta |", text)

    def test_extract_embedded_images(self) -> None:
        blobs = fe.extract_docx_images(self._make_docx(with_image=True))
        self.assertEqual(len(blobs), 1)
        self.assertEqual(chunking._raster_ext(blobs[0]), ".png")

    def test_no_images_returns_empty(self) -> None:
        self.assertEqual(fe.extract_docx_images(self._make_docx()), [])

    def test_describe_doc_images_appends_description(self) -> None:
        path = self._make_docx(with_image=True)
        seen: list[Path] = []

        def fake_vision(p: Path) -> str:
            seen.append(p)
            return "Схема: host-01 → подсистема Beta"

        block = chunking._describe_doc_images(path, fake_vision)
        self.assertIn("Изображения из документа", block)
        self.assertIn("Схема: host-01", block)
        self.assertEqual(len(seen), 1)  # одна картинка описана

    def test_describe_doc_images_no_vision_is_noop(self) -> None:
        self.assertEqual(chunking._describe_doc_images(self._make_docx(True), None), "")


class TestXlsxAllSheets(unittest.TestCase):
    def test_multiple_sheets_concatenated_with_label(self) -> None:
        with TemporaryDirectory() as d:
            p = Path(d) / "book.xlsx"
            with pd.ExcelWriter(str(p), engine="openpyxl") as w:
                pd.DataFrame({"x": [1, 2]}).to_excel(w, sheet_name="Лист1", index=False)
                pd.DataFrame({"x": [3, 4]}).to_excel(w, sheet_name="Лист2", index=False)
            df = fe._read_xlsx(p)
            self.assertIn("__sheet__", df.columns)
            self.assertEqual(set(df["__sheet__"]), {"Лист1", "Лист2"})
            self.assertEqual(len(df), 4)

    def test_single_sheet_unchanged(self) -> None:
        with TemporaryDirectory() as d:
            p = Path(d) / "one.xlsx"
            pd.DataFrame({"x": [1, 2]}).to_excel(str(p), index=False)
            df = fe._read_xlsx(p)
            self.assertNotIn("__sheet__", df.columns)
            self.assertEqual(len(df), 2)


if __name__ == "__main__":
    unittest.main()
